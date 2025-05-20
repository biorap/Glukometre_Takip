import os
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, timedelta, date
import locale
import csv
import sqlite3
import subprocess
import glob
import sys
from PIL import ImageTk, Image # Bu satırı ekleyin
from tkcalendar import Calendar
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Color
from openpyxl.utils import get_column_letter

# HBTC Kalite Kontrol Formu oluşturma için python-docx kütüphanesi
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_BREAK
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("Uyarı: Word formu oluşturma için 'python-docx' kütüphanesi bulunamadı. 'pip install python-docx' ile kurabilirsiniz.")

# Türkçe alfabetik sıralama için locale kütüphanesi
try:
    locale.setlocale(locale.LC_COLLATE, 'tr_TR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_COLLATE, 'turkish')
    except locale.Error:
        print("Uyarı: Türkçe yerel ayarları (tr_TR.UTF-8 veya turkish) bulunamadı. Sıralama varsayılan şekilde yapılacak.")

# Çıkış dizinleri
BACKUP_DIR = "Yedeklenmis Veriler" # .csv Yedekleme dizini
EXCEL_OUTPUT_DIR = "Excel'e Aktarılanlar" # Excel çıktıları için çıktı klasörü
WORD_FORMS_DIR = "HBTC Kalite Kontrol" # Word formları için çıktı klasörü
SABLONLAR_DIR = "Sablonlar"
HBTC_TEMPLATE_FILE = "HBTC_KALITE_KONTROL_FORMU.docx"
KALITE_KONTROL_SABLON_EXCEL = "Kalite_Kontrol_Verileri_Sablon.xlsx"
YUZDE_SAPMA_SABLON_EXCEL = "Yuzde_Sapma_Verileri_Sablon.xlsx"
CALENDAR_ICON_PATH = "Resources\\calendar.ico" # Takvim simgesi
APP_ICON_PATH = "Resources\\app_icon.ico" # Uygulama simgesi
# Veri dosyaları
DEVICES_TYPES_FILE = "DevicesTypes.txt" # Cihaz tipleri ve markalarını saklayan dosya
DEVICES_SERIALS_FILE = "DevicesSerials.txt" # Cihaz seri numaralarını saklayan dosya
BIRIMLER_FILE = "Birimler.txt" # Cihazın geldiği birimleri saklayan dosya
DB_FILE = "cihazlar.db" # Database dosyası
DB_SETTINGS_TABLE = "program_ayarlari" # Database dosyası içerisinde ayarlar tablosu

# Veri dosyaları bulunamazsa örnek veri dosyaları oluştur
SAMPLE_DEVICE_TYPES = ["GLUKOMETRE-BİOJECT PLUS", "ACCU-CHEK PERFORMA NANO", "BAYER CONTOUR NEXT", "İNOVATİF CİHAZ X", "ÖZEL MODEL Y"]
SAMPLE_DEVICE_SERIALS = ["BG709223125", "SN123456789", "SN987654321"]
SAMPLE_BIRIMLER = ["ACİL SERVİS", "DAHİLİYE POLİKLİNİĞİ", "YOĞUN BAKIM", "ÇOCUK SERVİSİ", "İNTANİYE", "ÜROLOJİ", "ACİL İZOLASYON"]

class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event=None):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25

        self.tooltip_window = tk.Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")

        label = tk.Label(self.tooltip_window, text=self.text, justify='left',
                        background="#ffffe0", relief='solid', borderwidth=1,
                        font=("tahoma", "10", "normal"))
        label.pack(ipadx=1)

    def hide_tooltip(self, event=None):
        if self.tooltip_window:
            self.tooltip_window.destroy()
        self.tooltip_window = None

class MainWindow:
    def __init__(self, master):
        self.master = master
        master.title("GLUKOMETRE TAKİP PROGRAMI")
        master.geometry("1300x700")
        master.resizable(True, True)
        master.iconbitmap(APP_ICON_PATH.replace("\\", "/"))

        self.renk_saat_bg = "#000000"
        self.renk_saat_fg = "#00ff00"

        self.style = ttk.Style()

        for dirname in [BACKUP_DIR, EXCEL_OUTPUT_DIR, WORD_FORMS_DIR, SABLONLAR_DIR]:
            if not os.path.exists(dirname):
                os.makedirs(dirname)

        self.init_device_db()
        self.init_settings_db()
        self.style = ttk.Style()

        self.main_frame = ttk.Frame(master)
        self.main_frame.pack(fill="both", expand=True)

        self.frm_sol_panel = ttk.Frame(self.main_frame, width=270, style="SolPanel.TFrame")
        self.frm_sol_panel.pack(side="left", fill="y", padx=(10,0), pady=10)
        self.frm_sol_panel.pack_propagate(False)

        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        self.frm_glukometre_genel = ttk.LabelFrame(self.frm_sol_panel, text="Glukometre Bilgileri", style="Glukometre.TLabelframe")
        self.frm_glukometre_genel.pack(padx=5, pady=5, fill="x", expand=False)

        ttk.Label(self.frm_glukometre_genel, text="Glukometrenin Geldiği Birim:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_birim = ttk.Combobox(self.frm_glukometre_genel, state="readonly")
        self.cmb_birim.pack(fill="x", padx=5, pady=(0,5))
        self.cmb_birim.bind("<<ComboboxSelected>>", self.on_birim_cihaz_secildi)

        frm_birim_buttons = ttk.Frame(self.frm_glukometre_genel)
        frm_birim_buttons.pack(fill="x", padx=5, pady=(0,10))

        # İkonları yükle
        try:
            # Betiğin bulunduğu dizini al
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # İkon dosyalarının tam yollarını oluştur
            plus_icon_path = os.path.join(script_dir, "Resources", "plus.ico")
            minus_icon_path = os.path.join(script_dir, "Resources", "minus.ico")

            if not os.path.exists(plus_icon_path):
                raise FileNotFoundError(f"Plus ikonu bulunamadı: {plus_icon_path}")
            if not os.path.exists(minus_icon_path):
                raise FileNotFoundError(f"Minus ikonu bulunamadı: {minus_icon_path}")

            # İkonları Pillow ile aç ve PhotoImage nesnesine dönüştür
            plus_image = Image.open(plus_icon_path)
            self.plus_icon = ImageTk.PhotoImage(plus_image)

            minus_image = Image.open(minus_icon_path)
            self.minus_icon = ImageTk.PhotoImage(minus_image)

        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.plus_icon = None
            self.minus_icon = None
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.plus_icon = None
            self.minus_icon = None

        # Butonları oluştur ve ikonları ekle
        if self.plus_icon:
            btn_birim_ekle = ttk.Button(frm_birim_buttons, text="Birim Ekle", image=self.plus_icon, compound=tk.LEFT, command=self.birim_ekle_pencere)
        else:
            btn_birim_ekle = ttk.Button(frm_birim_buttons, text="Birim Ekle", command=self.birim_ekle_pencere)
        self.btn_birim_ekle = btn_birim_ekle # self eklenerek sınıf değişkeni yapıldı
        btn_birim_ekle.pack(side="left", fill="x", expand=True, padx=(0,2))
        ToolTip(self.btn_birim_ekle, "Yeni bir birim eklemek için tıklayınız.")

        if self.minus_icon:
            btn_birim_sil = ttk.Button(frm_birim_buttons, text="Birim Sil", image=self.minus_icon, compound=tk.LEFT, command=self.birim_sil)
        else:
            btn_birim_sil = ttk.Button(frm_birim_buttons, text="Birim Sil", command=self.birim_sil)
        self.btn_birim_sil = btn_birim_sil # self eklenerek sınıf değişkeni yapıldı
        btn_birim_sil.pack(side="left", fill="x", expand=True, padx=(2,0))
        ToolTip(self.btn_birim_sil, "Seçili birimi silmek için tıklayınız.")

        ttk.Label(self.frm_glukometre_genel, text="Cihaz Tipi - Marka:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_device_type = ttk.Combobox(self.frm_glukometre_genel) # state="readonly" kaldırıldı
        self.cmb_device_type.pack(fill="x", padx=5, pady=(0,5))
        self.cmb_device_type.bind("<FocusOut>", self.on_device_type_entered) # Klavyeden giriş için
        self.cmb_device_type.bind("<Return>", self.on_device_type_entered) # Enter tuşu için

        frm_seri_no_container = ttk.Frame(self.frm_glukometre_genel)
        frm_seri_no_container.pack(fill="x", padx=0, pady=0)
        frm_seri_no_sol = ttk.Frame(frm_seri_no_container)
        frm_seri_no_sol.pack(side="left", fill="x", expand=True, padx=(5,2), pady=(0,5))
        ttk.Label(frm_seri_no_sol, text="Cihaz Seri No:").pack(fill="x")
        self.cmb_device_serial = ttk.Combobox(frm_seri_no_sol)
        self.cmb_device_serial.pack(fill="x")
        self.cmb_device_serial.bind("<Return>", self.on_seri_no_entered)
        self.cmb_device_serial.bind("<FocusOut>", self.on_seri_no_entered)

        frm_seri_no_sag = ttk.Frame(frm_seri_no_container)
        frm_seri_no_sag.pack(side="left", fill="x", expand=True, padx=(2,5), pady=(0,5))
        ttk.Label(frm_seri_no_sag, text="Son 4 Hane:").pack(fill="x")

        self.cmb_son4hane = ttk.Combobox(frm_seri_no_sag)
        self.cmb_son4hane.pack(fill="x")
        self.cmb_son4hane.bind("<<ComboboxSelected>>", self.on_son4hane_changed)
        self.cmb_son4hane.bind("<FocusOut>", self.on_son4hane_changed)
        self.cmb_son4hane.bind("<KeyRelease>", self.validate_son4hane_input)

        self.frm_radyo = ttk.LabelFrame(self.frm_sol_panel, text="Radyo", style="Radyo.TLabelframe")
        self.frm_radyo.pack(side="bottom", fill="x", padx=5, pady=(0,5)) # pack side bottom

        ttk.Label(self.frm_radyo, text="Radyo İstasyonu:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_radyo = ttk.Combobox(self.frm_radyo, state="readonly")
        self.radio_station_names, self.radio_station_map = self.load_radio_stations_from_db()
        ttk.Label(self.frm_radyo, text="Radyo İstasyonu:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_radyo['values'] = self.radio_station_names

        last_radio_station = self.load_setting("last_radio_station")
        if last_radio_station and last_radio_station in self.radio_station_names:
            self.cmb_radyo.set(last_radio_station)
        elif self.radio_station_names: # Eğer liste boş değilse
            self.cmb_radyo.current(0)

        self.load_radio_stations_from_db() # Çağrı şekli değiştirildi
        self.cmb_radyo.pack(fill="x", padx=5, pady=(0,5)) # Şimdi combobox'ı paketleyin
        last_radio_station = self.load_setting("last_radio_station")
        if last_radio_station and last_radio_station in self.radio_station_names:
            self.cmb_radyo.set(last_radio_station)
        elif self.radio_station_names:
            self.cmb_radyo.current(0)

        frm_radyo_controls = ttk.Frame(self.frm_radyo)
        frm_radyo_controls.pack(fill="x", padx=5, pady=(0,10))

        # İkonları yükle
        try:
            # Betiğin bulunduğu dizini al
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # İkon dosyalarının tam yollarını oluştur
            play_icon_path = os.path.join(script_dir, "Resources", "play.ico")
            stop_icon_path = os.path.join(script_dir, "Resources", "stop.ico")

            if not os.path.exists(play_icon_path):
                raise FileNotFoundError(f"Play ikonu bulunamadı: {play_icon_path}")
            if not os.path.exists(stop_icon_path):
                raise FileNotFoundError(f"Stop ikonu bulunamadı: {stop_icon_path}")

            # İkonları Pillow ile aç ve PhotoImage nesnesine dönüştür
            play_image = Image.open(play_icon_path)
            self.play_icon = ImageTk.PhotoImage(play_image)

            stop_image = Image.open(stop_icon_path)
            self.stop_icon = ImageTk.PhotoImage(stop_image)

        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.play_icon = None
            self.stop_icon = None
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.play_icon = None
            self.stop_icon = None

        # Butonları oluştur ve ikonları ekle
        # Bu kısım aynı kalabilir, çünkü self.play_icon ve self.stop_icon artık PhotoImage nesneleri (veya None)
        if self.play_icon:
            self.btn_radyo_play = ttk.Button(frm_radyo_controls, text="Play", image=self.play_icon, compound=tk.LEFT, command=self.play_radio_command)
        else:
            self.btn_radyo_play = ttk.Button(frm_radyo_controls, text="Play", command=self.play_radio_command)

        self.btn_radyo_play.grid(row=0, column=0, sticky="ew", padx=(0,1))

        if self.stop_icon:
            self.btn_radyo_stop = ttk.Button(frm_radyo_controls, text="Stop", image=self.stop_icon, compound=tk.LEFT, command=self.stop_radio)
        else:
            self.btn_radyo_stop = ttk.Button(frm_radyo_controls, text="Stop", command=self.stop_radio)

        self.btn_radyo_stop.grid(row=0, column=1, sticky="ew", padx=(1,0))

        self.radio_volume = tk.IntVar(value=50)
        last_volume = self.load_setting("last_radio_volume", "50")
        try:
            self.radio_volume.set(int(last_volume))
        except ValueError:
            self.radio_volume.set(50)

        self.volume_slider = ttk.Scale(frm_radyo_controls, from_=0, to=100, orient=tk.HORIZONTAL, variable=self.radio_volume, command=self.on_volume_change)
        self.volume_slider.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(5,0))
        ToolTip(self.volume_slider, "Ses Seviyesi")
        
        # "Sesi Kıs" butonu
        try:
            # Betiğin bulunduğu dizini al
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # İkon dosyalarının tam yollarını oluştur
            mute_icon_path = os.path.join(script_dir, "Resources", "mute.ico")
            loud_icon_path = os.path.join(script_dir, "Resources", "loud.ico")

            if not os.path.exists(mute_icon_path):
                raise FileNotFoundError(f"Mute ikonu bulunamadı: {mute_icon_path}")
            if not os.path.exists(loud_icon_path):
                raise FileNotFoundError(f"Loud ikonu bulunamadı: {loud_icon_path}")

            # İkonları Pillow ile aç ve PhotoImage nesnesine dönüştür
            mute_image = Image.open(mute_icon_path)
            self.mute_icon = ImageTk.PhotoImage(mute_image)

            loud_image = Image.open(loud_icon_path)
            self.loud_icon = ImageTk.PhotoImage(loud_image)

            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", image=self.mute_icon, compound=tk.LEFT, command=self.toggle_mute_sound)
        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.mute_icon = None
            self.loud_icon = None
            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", command=self.toggle_mute_sound)
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.mute_icon = None
            self.loud_icon = None
            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", command=self.toggle_mute_sound)

        self.btn_mute_sound.grid(row=1, column=2, sticky="ew", padx=(5,0), pady=(5,0))
        ToolTip(self.btn_mute_sound, "Ses seviyesini sıfırla/eski haline getir")
        self.lbl_volume_percent = ttk.Label(frm_radyo_controls, text=f"Ses Seviyesi: {self.radio_volume.get()}%")
        self.lbl_volume_percent.grid(row=2, column=0, columnspan=3, sticky="n") # columnspan 3 yapıldı

        frm_radyo_controls.columnconfigure(0, weight=1)
        frm_radyo_controls.columnconfigure(1, weight=1)
        frm_radyo_controls.columnconfigure(2, weight=0) # Sesi Kıs butonu için

        self.radio_process = None
        self.update_radio_button_states()
        self.cmb_radyo.bind("<<ComboboxSelected>>", self.on_radio_station_selected)

        # Mute durumunu takip etmek için bir değişken
        self.is_muted = False
        self.previous_volume = self.radio_volume.get() # Başlangıçta mevcut ses seviyesini sakla

        # Dijital Saat en altta olacak
        self.lbl_dijital_saat = tk.Label(self.frm_sol_panel, text="", font=("Arial", 16, "bold"),
                                          background=self.renk_saat_bg, foreground=self.renk_saat_fg,
                                          relief="sunken", borderwidth=2)
        self.lbl_dijital_saat.pack(side="bottom", fill="x", padx=5, pady=(0,5))

        # İşlem Label'ı
        self.lbl_islem_durumu = ttk.Label(self.frm_sol_panel, text="Veri İşleniyor. Lütfen Bekleyiniz...",
                                          font=("Arial", 12, "italic"), background="black", foreground="white",)
        self.islem_label_visible = False
        self.blinking_after_id = None

        self.status_bar = tk.Label(self.master, text="", relief="sunken", anchor="w", font=("Arial", 10))
        self.status_bar.pack(side="bottom", fill="x")

        self.guncelle_dijital_saat()

        self.load_data_from_files_to_db()
        self.load_initial_data()
        # Combobox seçim değişikliklerini ayarlara kaydet
        self.cmb_device_type.bind("<<ComboboxSelected>>", self.on_device_type_selected)
        self.cmb_device_serial.bind("<<ComboboxSelected>>", self.on_device_serial_selected)
        self.on_birim_cihaz_secildi()

        self.measurement_no_kalite = 1
        self.measurement_no_yuzde = 1
        self.editing_entry = None
        self.l_entry_tooltips = {}
        self.tables_cleared_this_session = False

        self.create_tabs()
        self.create_menu()
        self.load_data_from_csv()
        self.update_status_bar()
        master.protocol("WM_DELETE_WINDOW", self.on_closing)

        # Cihaz kaydı çakışma kontrolünü tablardaki textbox kutularına odaklanınca yap
        self.txt_l1.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_l2.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_l3.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_glukometre_yuzde.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_lab_yuzde.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)

        self.cakisma_uyarildi = False  # Uyarı gösterildi mi kontrolü

    def start_islem_label(self):
        if not self.islem_label_visible:
            self.lbl_islem_durumu.pack(side="bottom", fill="x", padx=5, pady=(5,0))
            self.islem_label_visible = True
        self._blink_islem_label()

    def stop_islem_label(self):
        if self.islem_label_visible:
            self.lbl_islem_durumu.pack_forget()
            self.islem_label_visible = False
            if self.blinking_after_id:
                self.master.after_cancel(self.blinking_after_id)
                self.blinking_after_id = None
            self.lbl_islem_durumu.config(foreground="blue") # Rengi normale döndür

    def _blink_islem_label(self):
        if self.islem_label_visible:
            current_color = self.lbl_islem_durumu.cget("foreground")
            new_color = "blue" if str(current_color) == "white" else "white"
            self.lbl_islem_durumu.config(foreground=new_color)
            self.blinking_after_id = self.master.after(500, self._blink_islem_label)

    def init_settings_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS {DB_SETTINGS_TABLE} (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)
        conn.commit()
        conn.close()

    def save_setting(self, key, value):
        try:
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            cursor.execute(f"INSERT OR REPLACE INTO {DB_SETTINGS_TABLE} (key, value) VALUES (?, ?)", (key, value))
            conn.commit()
            conn.close()
        except sqlite3.Error as e:
            print(f"DB Ayar kaydetme hatası ({key}): {e}")

    def load_setting(self, key, default=None):
        try:
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            cursor.execute(f"SELECT value FROM {DB_SETTINGS_TABLE} WHERE key = ?", (key,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else default
        except sqlite3.Error as e:
            print(f"DB Ayar yükleme hatası ({key}): {e}")
            return default

    def on_radio_station_selected(self, event=None):
        # Called when a new radio station is selected from the combobox
        if self.radio_process and self.radio_process.poll() is None: # If radio is currently playing
            self.stop_radio()
            # Wait a bit for ffplay to stop, then play new station with current volume
            self.master.after(200, self.play_radio_command)

    def on_volume_change(self, value_str):
        volume_level = int(float(value_str))
        self.radio_volume.set(volume_level) # Ensure IntVar is updated
        self.lbl_volume_percent.config(text=f"Ses Seviyesi: {volume_level}%")
        self.save_setting("last_radio_volume", str(volume_level))

        if self.radio_process and self.radio_process.poll() is None: # If radio is currently playing
            self.stop_radio()
            # Wait a bit for ffplay to stop, then restart with new volume
            self.master.after(200, lambda: self.play_radio(volume_level=volume_level))
            
    def toggle_mute_sound(self):
        if not hasattr(self, 'loud_icon') or not hasattr(self, 'mute_icon'):
            messagebox.showerror("Hata", "Loud/Mute ikonları yüklenemedi!")
            return

        if not self.is_muted:
            # Sesi kıs
            self.previous_volume = self.radio_volume.get()
            self.radio_volume.set(0)
            self.volume_slider.set(0)
            self.lbl_volume_percent.config(text=f"Ses Seviyesi: 0%")
            self.save_setting("last_radio_volume", "0")

            # Butonun görünümünü değiştir
            self.btn_mute_sound.config(image=self.loud_icon, text="Sesi Aç")
        else:
            # Sesi aç
            self.radio_volume.set(self.previous_volume)
            self.volume_slider.set(self.previous_volume)
            self.lbl_volume_percent.config(text=f"Ses Seviyesi: {self.previous_volume}%")
            self.save_setting("last_radio_volume", str(self.previous_volume))

            # Butonun görünümünü değiştir
            self.btn_mute_sound.config(image=self.mute_icon, text="Sesi Kıs")

        self.is_muted = not self.is_muted
        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()
            self.master.after(200, lambda: self.play_radio(volume_level=self.radio_volume.get()))
        else: # Radyo çalmıyorsa, sadece ses seviyesini sıfırla
            self.update_radio_button_states(playing=False)

    def mute_sound(self):
        self.radio_volume.set(0)
        self.volume_slider.set(0) # Slider'ı da güncelle
        self.lbl_volume_percent.config(text=f"Ses Seviyesi: 0%")
        self.save_setting("last_radio_volume", "0")
        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()
        else: # Radyo çalmıyorsa, sadece ses seviyesini sıfırla
            self.update_radio_button_states(playing=False)

    def update_radio_button_states(self, playing=False):
        if playing:
            self.btn_radyo_play.config(state=tk.DISABLED)
            self.btn_radyo_stop.config(state=tk.NORMAL)
        else:
            self.btn_radyo_play.config(state=tk.NORMAL)
            self.btn_radyo_stop.config(state=tk.DISABLED)

    def update_status_bar(self):
        current_date = datetime.now().strftime("%d.%m.%Y")
        kalite_count = len(self.tree_kalite.get_children()) if hasattr(self, 'tree_kalite') else 0
        yuzde_count = len(self.tree_yuzde.get_children()) if hasattr(self, 'tree_yuzde') else 0
        status_text = f"Tarih: {current_date} | Kalite Kontrol Ölçümleri: {kalite_count} | Yüzde Sapma Ölçümleri: {yuzde_count}"
        self.status_bar.config(text=status_text)

    def init_device_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()

        # Birimler Tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS birimler (
                birim_adi TEXT PRIMARY KEY
            )
        """)

        # Cihaz Tipleri Tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cihaz_tipleri (
                cihaz_tipi TEXT PRIMARY KEY
            )
        """)

        # Cihaz Serileri Tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cihaz_serileri (
                cihaz_seri TEXT PRIMARY KEY
            )
        """)

        # Radyolar Tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS radyolar (
                radyo_adi TEXT PRIMARY KEY,
                radyo_url TEXT NOT NULL
            )
        """)

        # Cihaz Kayıtları Tablosu (Zaten Mevcut)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cihaz_kayitlari (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                birim_adi TEXT NOT NULL,
                cihaz_tipi TEXT NOT NULL,
                cihaz_seri TEXT NOT NULL,
                son_4_hane TEXT NOT NULL,
                UNIQUE (cihaz_tipi, cihaz_seri, son_4_hane)
            )
        """)

        conn.commit()
        conn.close()

    def load_initial_data(self):
        self.load_birimler_from_db()
        self.load_device_types_from_db()
        self.load_device_serials_from_db()
        self.load_radio_stations_from_db()

    def load_birimler_from_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT birim_adi FROM birimler")
        results = cursor.fetchall()
        self.birimler = [row[0] for row in results]
        conn.close()
        self.cmb_birim['values'] = self.birimler
        if self.birimler:
            self.cmb_birim.current(0)

    def load_device_types_from_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT cihaz_tipi FROM cihaz_tipleri")
        results = cursor.fetchall()
        self.device_types = sorted([row[0] for row in results], key=locale.strxfrm)
        conn.close()
        self.cmb_device_type['values'] = self.device_types

        last_type = self.load_setting("last_selected_device_type")
        if last_type and last_type in self.device_types:
            self.cmb_device_type.set(last_type)
        else:
            self.cmb_device_type.set("")


    def load_device_serials_from_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT cihaz_seri FROM cihaz_serileri")
        results = cursor.fetchall()
        self.device_serials = sorted([row[0] for row in results], key=locale.strxfrm)
        conn.close()
        self.cmb_device_serial['values'] = self.device_serials

        # Otomatik seçim kaldırıldı, ayarlardan yükle
        last_serial = self.load_setting("last_selected_serial_no")
        if last_serial and last_serial in self.device_serials:
            self.cmb_device_serial.set(last_serial)
        else:
            self.cmb_device_serial.set("")

    # GlukometreTakip.py dosyasında, MainWindow sınıfı içinde
    def load_radio_stations_from_db(self):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT radyo_adi, radyo_url FROM radyolar")
        results = cursor.fetchall()
        # Yerel değişkenlere ata
        local_radio_station_names = [row[0] for row in results]
        local_radio_station_map = {row[0]: row[1] for row in results}
        conn.close()

        # Sınıf özelliklerini de ayarla (isteğe bağlı, __init__ içinde de yapılabilir)
        self.radio_station_names = local_radio_station_names
        self.radio_station_map = local_radio_station_map

        return local_radio_station_names, local_radio_station_map # Veriyi DÖNDÜR

    def load_data_from_files_to_db(self):
        # Birimleri yükle
        if os.path.exists(BIRIMLER_FILE):
            birimler = self.load_file_lines(BIRIMLER_FILE, sort_data=True)
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            for birim in birimler:
                cursor.execute("INSERT OR IGNORE INTO birimler (birim_adi) VALUES (?)", (birim,))
            conn.commit()
            conn.close()

        # Cihaz tiplerini yükle
        if os.path.exists(DEVICES_TYPES_FILE):
            cihaz_tipleri = self.load_file_lines(DEVICES_TYPES_FILE, sort_data=True)
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            for cihaz_tipi in cihaz_tipleri:
                cursor.execute("INSERT OR IGNORE INTO cihaz_tipleri (cihaz_tipi) VALUES (?)", (cihaz_tipi,))
            conn.commit()
            conn.close()

        # Cihaz serilerini yükle
        if os.path.exists(DEVICES_SERIALS_FILE):
            cihaz_serileri = self.load_file_lines(DEVICES_SERIALS_FILE, sort_data=False)
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            for cihaz_seri in cihaz_serileri:
                cursor.execute("INSERT OR IGNORE INTO cihaz_serileri (cihaz_seri) VALUES (?)", (cihaz_seri,))
            conn.commit()
            conn.close()

        # Radyo istasyonlarını yükle
        if os.path.exists("RadioStationsFFMPEG.txt"):
            radio_stations, radio_map = self.load_radio_stations_from_file()
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            for name, url in radio_map.items():
                cursor.execute("INSERT OR IGNORE INTO radyolar (radyo_adi, radyo_url) VALUES (?, ?)", (name, url))
            conn.commit()
            conn.close()

    def get_son4hane_for_device(self, birim, tip, seri):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT son_4_hane FROM cihaz_kayitlari
            WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ?
            ORDER BY id DESC LIMIT 1
        """, (birim, tip, seri))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else ""

    def get_son4hane_list_for_device(self, birim, tip, seri):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT DISTINCT son_4_hane FROM cihaz_kayitlari
            WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ?
        """, (birim, tip, seri))
        results = cursor.fetchall()
        conn.close()
        return [row[0] for row in results] if results else []

    def add_or_update_device_assignment(self, birim, tip, seri, son4):
        if not (birim and tip and seri and son4 and len(son4)==4):
            messagebox.showerror("Hata", "Lütfen tüm alanları doldurun ve Seri numarasının son 4 hanesini kontrol edin.")
            return False

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        try:
            # Aynı birim, seri, son4hane başka bir cihaz tipine kayıtlı mı?
            cursor.execute("""
                SELECT cihaz_tipi FROM cihaz_kayitlari
                WHERE birim_adi = ? AND cihaz_seri = ? AND son_4_hane = ? AND cihaz_tipi != ?
            """, (birim, seri, son4, tip))
            existing_assignment = cursor.fetchone()
            if existing_assignment:
                messagebox.showerror(
                    "Kayıt Hatası",
                    f" '{existing_assignment[0]}' cihazı, '{seri}{son4}' seri numarasıyla zaten '{birim}' biriminde kayıtlı.\n"
                    "Cihaz Tipi, Seri No ve Son 4 Haneyi verilerini tekrar kontrol edin."
                )
                return False

            # Eski kontrol: Aynı cihaz başka bir birime kayıtlı mı?
            cursor.execute("""
                SELECT birim_adi FROM cihaz_kayitlari
                WHERE cihaz_tipi = ? AND cihaz_seri = ? AND son_4_hane = ? AND birim_adi != ?
            """, (tip, seri, son4, birim))
            existing_assignment_birim = cursor.fetchone()
            if existing_assignment_birim:
                messagebox.showerror(
                    "Kayıt Hatası",
                    f"Bu Seri Numaralı cihaz ({seri}{son4}) zaten '{existing_assignment_birim[0]}' birimine kayıtlı.\n"
                    "Girdiğiniz Son 4 Haneyi tekrar kontrol ediniz."
                )
                return False

            cursor.execute("""
                INSERT OR IGNORE INTO cihaz_kayitlari (birim_adi, cihaz_tipi, cihaz_seri, son_4_hane)
                VALUES (?, ?, ?, ?)
            """, (birim, tip, seri, son4))
            conn.commit()
            self.on_birim_cihaz_secildi() # Combobox'ı güncelle
            return True
        except sqlite3.Error as e:
            messagebox.showerror("Veritabanı Hatası", f"Cihaz kaydı sırasında hata: {e}")
            return False
        finally:
            conn.close()

    def kontrol_cihaz_kayit_cakisma(self, event=None):
        if getattr(self, "cakisma_uyarildi", False):
            return  # Zaten uyarı gösterildi, tekrar gösterme

        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        son4 = self.cmb_son4hane.get().strip().upper()
        if not (birim and tip and seri and son4 and len(son4) == 4):
            return  # Tüm alanlar dolu değilse kontrol etme

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT birim_adi, cihaz_tipi FROM cihaz_kayitlari
                WHERE cihaz_seri = ? AND son_4_hane = ? AND (birim_adi != ? OR cihaz_tipi != ?)
            """, (seri, son4, birim, tip))
            result = cursor.fetchone()
            if result:
                self.cakisma_uyarildi = True
                messagebox.showerror(
                    "Kayıt Çakışması",
                    f"Bu Seri Numaralı cihaz ({seri}{son4}) zaten '{result[0]}' biriminde ve '{result[1]}' cihaz tipiyle kayıtlı.\n"
                    "Cihaz Tipi, Seri No, Son 4 Hane ve Birim bilgilerini kontrol ediniz."
                )
                self.master.after(100, lambda: setattr(self, "cakisma_uyarildi", False))
                self.master.focus()
                return "break"
        finally:
            conn.close()


    def guncelle_dijital_saat(self):
        simdiki_zaman = datetime.now().strftime("%H:%M:%S")
        self.lbl_dijital_saat.config(text=simdiki_zaman)
        self.master.after(1000, self.guncelle_dijital_saat)

    def check_device_availability(self, birim, tip, seri, son4):
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT birim_adi FROM cihaz_kayitlari
            WHERE cihaz_tipi = ? AND cihaz_seri = ? AND son_4_hane = ? AND birim_adi != ?
        """, (tip, seri, son4, birim))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else None

    def on_birim_cihaz_secildi(self, event=None):
        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        if birim and tip and seri:
            son4_list = self.get_son4hane_list_for_device(birim, tip, seri)
            self.cmb_son4hane['values'] = son4_list
            if son4_list:
                self.cmb_son4hane.set(son4_list[0])
            else:
                self.cmb_son4hane.set("")
            self.validate_son4hane_input()
        else:
            self.cmb_son4hane['values'] = []
            self.cmb_son4hane.set("")
            self.validate_son4hane_input()

    def on_son4hane_changed(self, event=None):
        self.validate_son4hane_input()

    def on_seri_no_entered(self, event=None):
        new_serial = self.cmb_device_serial.get().strip().upper()
        if new_serial and new_serial not in self.device_serials:
            try:
                conn = sqlite3.connect(DB_FILE)
                cursor = conn.cursor()
                cursor.execute("INSERT OR IGNORE INTO cihaz_serileri (cihaz_seri) VALUES (?)", (new_serial,))
                conn.commit()
                conn.close()
                self.load_device_serials_from_db()
                self.cmb_device_serial.set(new_serial)
                self.on_birim_cihaz_secildi()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz seri no eklenirken hata: {e}")

    def on_device_type_entered(self, event=None):
        new_device_type = self.cmb_device_type.get().strip().upper()
        if new_device_type and new_device_type not in self.device_types:
            try:
                conn = sqlite3.connect(DB_FILE)
                cursor = conn.cursor()
                cursor.execute("INSERT OR IGNORE INTO cihaz_tipleri (cihaz_tipi) VALUES (?)", (new_device_type,))
                conn.commit()
                conn.close()
                self.load_device_types_from_db()
                self.cmb_device_type.set(new_device_type)
                self.on_birim_cihaz_secildi()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz tipi eklenirken hata: {e}")

    def on_device_type_selected(self, event=None):
        selected_type = self.cmb_device_type.get()
        self.save_setting("last_selected_device_type", selected_type)

    def on_device_serial_selected(self, event=None):
        selected_serial = self.cmb_device_serial.get()
        self.save_setting("last_selected_serial_no", selected_serial)

    def validate_son4hane_input(self, event=None):
        content = self.cmb_son4hane.get().upper()
        new_content = "".join(filter(str.isalnum, content))[:4]
        if self.cmb_son4hane.get() != new_content:
            self.cmb_son4hane.set(new_content)

        if len(new_content) == 4:
            self.cmb_son4hane.config(style="Valid.TCombobox")
            if hasattr(self, 'son4hane_tooltip') and self.son4hane_tooltip.tooltip_window:
                self.son4hane_tooltip.hide_tooltip()
        elif len(new_content) > 0:
            self.cmb_son4hane.config(style="Invalid.TCombobox")
            if not hasattr(self, 'son4hane_tooltip') or self.son4hane_tooltip.widget != self.cmb_son4hane:
                self.son4hane_tooltip = ToolTip(self.cmb_son4hane, "Cihazın seri numarasının son 4 hanesini giriniz ya da varsa listeden seçiniz.")
        else:
            self.cmb_son4hane.config(style="TCombobox")
            if hasattr(self, 'son4hane_tooltip') and self.son4hane_tooltip.tooltip_window:
                self.son4hane_tooltip.hide_tooltip()
        return True

    def load_file_lines(self, filename, sample_data=None, sort_data=True):
        if not os.path.exists(filename):
            with open(filename, "w", encoding="utf-8-sig") as f:
                if sample_data:
                    for line in sample_data: f.write(line.upper() + "\n")
            lines = [line.upper() for line in sample_data] if sample_data else []
        else:
            with open(filename, "r", encoding="utf-8-sig") as f:
                lines = [line.strip().upper() for line in f if line.strip()]

        lines = list(dict.fromkeys(lines))
        if sort_data:
            try:
                return sorted(lines, key=locale.strxfrm)
            except NameError:
                return sorted(lines)
        return lines

    def save_birimler(self):
        try:
            sorted_birimler = sorted(self.birimler, key=locale.strxfrm)
        except NameError:
            sorted_birimler = sorted(self.birimler)
        with open(BIRIMLER_FILE, "w", encoding="utf-8-sig") as f:
            for birim in sorted_birimler: f.write(birim + "\n")
        self.cmb_birim['values'] = sorted_birimler

    def create_tabs(self):
        self.tab1 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text="Kalite Kontrol Ölçümleri")
        self.tab1.grid_columnconfigure(0, weight=1)
        self.tab1.grid_rowconfigure(2, weight=1)

        lbl_tablo_baslik_kalite = ttk.Label(self.tab1, text="KALİTE KONTROL ÖLÇÜMLERİ", font=("Tahoma", 14, "bold"))
        lbl_tablo_baslik_kalite.grid(row=0, column=0, pady=(10,5), sticky="n")
        self.tab1.grid_rowconfigure(0, weight=0)

        frm_kalite_input = ttk.LabelFrame(self.tab1, text="Kalite Kontrol Değer Girişi", style="KaliteInput.TLabelframe")
        frm_kalite_input.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        vcmd_l1 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 36, 108)), '%P', '%W')
        vcmd_l2 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 144, 216)), '%P', '%W')
        vcmd_l3 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 252, 396)), '%P', '%W')
        lbl_l1 = ttk.Label(frm_kalite_input, text="L1 Ölçümü:", foreground="#FFA200")
        lbl_l1.grid(row=0, column=0, padx=5, pady=5)
        self.txt_l1 = ttk.Entry(frm_kalite_input, width=10, name="l1_entry",
                              validate="key", validatecommand=vcmd_l1)
        self.txt_l1.grid(row=0, column=1, padx=5, pady=5)
        self.l_entry_tooltips["l1_entry"] = ToolTip(self.txt_l1, "Seviye 1 (L1) Değeri 36-108 arası olmalı")

        lbl_l2 = ttk.Label(frm_kalite_input, text="L2 Ölçümü:", foreground="#0000FF")
        lbl_l2.grid(row=0, column=2, padx=5, pady=5)
        self.txt_l2 = ttk.Entry(frm_kalite_input, width=10, name="l2_entry",
                              validate="key", validatecommand=vcmd_l2)
        self.txt_l2.grid(row=0, column=3, padx=5, pady=5)
        self.l_entry_tooltips["l2_entry"] = ToolTip(self.txt_l2, "Seviye 2 (L2) Değeri  144-216 arası olmalı")

        lbl_l3 = ttk.Label(frm_kalite_input, text="L3 Ölçümü:", foreground="#FF0000")
        lbl_l3.grid(row=0, column=4, padx=5, pady=5)
        self.txt_l3 = ttk.Entry(frm_kalite_input, width=10, name="l3_entry",
                              validate="key", validatecommand=vcmd_l3)
        self.txt_l3.grid(row=0, column=5, padx=5, pady=5)
        self.l_entry_tooltips["l3_entry"] = ToolTip(self.txt_l3, "Seviye 3 (L3) Değeri 252-396 arası olmalı")


        ttk.Button(frm_kalite_input, text="Tabloya Aktar", command=self.tabloya_aktar_kalite).grid(row=0, column=6, padx=20, pady=5)

        columns_kalite = (
            "No", "Tarih", "Cihaz Tipi - Marka", "Cihaz Seri No", "L1", "L2", "L3",
            "Glukometrenin Geldiği Birim", "Bir Sonraki Gelinecek Tarih"
        )
        self.tree_kalite = ttk.Treeview(self.tab1, columns=columns_kalite, show="headings", selectmode="extended")
        self.tree_kalite.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

        vsb_kalite = ttk.Scrollbar(self.tab1, orient="vertical", command=self.tree_kalite.yview)
        vsb_kalite.grid(row=2, column=1, sticky='ns', pady=(10,10))
        hsb_kalite = ttk.Scrollbar(self.tab1, orient="horizontal", command=self.tree_kalite.xview)
        hsb_kalite.grid(row=3, column=0, sticky='ew', padx=(10,10))
        self.tree_kalite.configure(yscrollcommand=vsb_kalite.set, xscrollcommand=hsb_kalite.set)

        widths_kalite = [20, 50, 150, 80, 30, 30, 30, 200, 140] # Sütun genişlikleri
        for col, w in zip(columns_kalite, widths_kalite):
            self.tree_kalite.heading(col, text=col, command=lambda c=col: self.treeview_sort_column(self.tree_kalite, c, False))
            self.tree_kalite.column(col, width=w, anchor=tk.CENTER, minwidth=w)

        self.tree_kalite.bind("<Double-1>", lambda event: self.on_double_click(event, self.tree_kalite))
        self.tree_kalite.bind("<Button-3>", lambda event: self.show_context_menu(event, self.tree_kalite))
        self.tree_kalite.bind("<Delete>", lambda event: self.satir_sil(self.tree_kalite, True))

        self.context_menu_kalite = tk.Menu(self.tab1, tearoff=0)
        self.context_menu_kalite.add_command(label="Satırı Sil", command=lambda: self.satir_sil(self.tree_kalite, True))

        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab2, text="Yüzde Sapma Ölçümleri")
        self.tab2.grid_columnconfigure(0, weight=1)
        self.tab2.grid_rowconfigure(2, weight=1)

        lbl_tablo_baslik_yuzde = ttk.Label(self.tab2, text="YÜZDE SAPMA ÖLÇÜMLERİ", font=("Tahoma", 14, "bold"))
        lbl_tablo_baslik_yuzde.grid(row=0, column=0, pady=(10,5), sticky="n")
        self.tab2.grid_rowconfigure(0, weight=0)


        frm_yuzde_input = ttk.LabelFrame(self.tab2, text="Yüzde Sapma Değer Girişi", style="YuzdeInput.TLabelframe")
        frm_yuzde_input.grid(row=1, column=0, padx=10, pady=5, sticky="ew")

        ttk.Label(frm_yuzde_input, text="Glukometre Ölçümü:").grid(row=0, column=0, padx=5, pady=5)
        self.txt_glukometre_yuzde = ttk.Entry(frm_yuzde_input, width=10)
        self.txt_glukometre_yuzde.grid(row=0, column=1, padx=5, pady=5)
        ToolTip(self.txt_glukometre_yuzde, "Gelen kan numunesinin Glukometre cihazında ölçtüğünüz ölçüm değerini giriniz.")    
        
        ttk.Label(frm_yuzde_input, text="Laboratuvar Ölçümü:").grid(row=0, column=2, padx=5, pady=5)
        self.txt_lab_yuzde = ttk.Entry(frm_yuzde_input, width=10)
        self.txt_lab_yuzde.grid(row=0, column=3, padx=5, pady=5)
        ToolTip(self.txt_lab_yuzde, "Gelen kan numunesinin Laboratuar Otoanalizör cihazında ölçülen ölçüm değerini giriniz.")    

        ttk.Button(frm_yuzde_input, text="Hesapla ve Tabloya Aktar", command=self.yuzde_sapma_hesapla_ve_aktar).grid(row=0, column=4, padx=20, pady=5)

        columns_yuzde = ("No", "Tarih", "Cihaz Tipi-Marka", "Cihaz Seri No", "Glukometre Ölçümü", "Lab. Ölçümü", "Yüzde Sapma", "Glukometrenin Geldiği Birim", "Bir Sonraki Gelinecek Tarih")
        self.tree_yuzde = ttk.Treeview(self.tab2, columns=columns_yuzde, show="headings", selectmode="extended")
        self.tree_yuzde.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

        vsb_yuzde = ttk.Scrollbar(self.tab2, orient="vertical", command=self.tree_yuzde.yview)
        vsb_yuzde.grid(row=2, column=1, sticky='ns', pady=(10,10))
        hsb_yuzde = ttk.Scrollbar(self.tab2, orient="horizontal", command=self.tree_yuzde.xview)
        hsb_yuzde.grid(row=3, column=0, sticky='ew', padx=(10,10))
        self.tree_yuzde.configure(yscrollcommand=vsb_yuzde.set, xscrollcommand=hsb_yuzde.set)
        self.tree_yuzde.tag_configure('high_deviation_tree', background='red', foreground='white')

        widths_yuzde = [20, 50, 155, 80, 100, 70, 70, 150, 120]
        for col, w in zip(columns_yuzde, widths_yuzde):
            self.tree_yuzde.heading(col, text=col, command=lambda c=col: self.treeview_sort_column(self.tree_yuzde, c, False))
            self.tree_yuzde.column(col, width=w, anchor=tk.CENTER, minwidth=w)

        self.tree_yuzde.bind("<Double-1>", lambda event: self.on_double_click(event, self.tree_yuzde))
        self.tree_yuzde.bind("<Button-3>", lambda event: self.show_context_menu(event, self.tree_yuzde))
        self.tree_yuzde.bind("<Delete>", lambda event: self.satir_sil(self.tree_yuzde, False))

        self.context_menu_yuzde = tk.Menu(self.tab2, tearoff=0)
        self.context_menu_yuzde.add_command(label="Satırı Sil", command=lambda: self.satir_sil(self.tree_yuzde, False))

    def validate_l_entry(self, P, widget_name, min_val, max_val):
        widget = self.master.nametowidget(widget_name)
        if P == "":
            widget.config(style="TEntry")
            if widget_name in self.l_entry_tooltips and self.l_entry_tooltips[widget_name].tooltip_window:
                self.l_entry_tooltips[widget_name].hide_tooltip()
            return True
        try:
            value = int(P)
            if min_val <= value <= max_val:
                widget.config(style="Valid.TEntry")
                if widget_name in self.l_entry_tooltips and self.l_entry_tooltips[widget_name].tooltip_window:
                    self.l_entry_tooltips[widget_name].hide_tooltip()
                return True
            else:
                widget.config(style="Invalid.TEntry")
                if widget_name in self.l_entry_tooltips:
                    self.l_entry_tooltips[widget_name].show_tooltip()
                return True
        except ValueError:
            widget.config(style="Invalid.TEntry")
            if widget_name in self.l_entry_tooltips:
                self.l_entry_tooltips[widget_name].show_tooltip()
            return True

    def check_l_entries_valid_for_aktar(self):
        valid_l1, valid_l2, valid_l3 = True, True, True
        try:
            l1_val = int(self.txt_l1.get())
            if not (36 <= l1_val <= 108): valid_l1 = False
        except ValueError: valid_l1 = False
        if not self.txt_l1.get(): valid_l1 = False

        try:
            l2_val = int(self.txt_l2.get())
            if not (144 <= l2_val <= 216): valid_l2 = False
        except ValueError: valid_l2 = False
        if not self.txt_l2.get(): valid_l2 = False

        try:
            l3_val = int(self.txt_l3.get())
            if not (252 <= l3_val <= 396): valid_l3 = False
        except ValueError: valid_l3 = False
        if not self.txt_l3.get(): valid_l3 = False

        if not (valid_l1 and valid_l2 and valid_l3):
            messagebox.showerror("Hata", "L1, L2, L3 değerlerini kontrol edin ve referans aralıklarında değerler girin.")
            return None
        return l1_val, l2_val, l3_val

    def create_menu(self):
        menubar = tk.Menu(self.master)
        menu_dosya = tk.Menu(menubar, tearoff=0)
        menu_dosya.add_command(label="Verileri Yedekten Geri Yükle...", command=self.manuel_yedek_yukle)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Tüm Tabloları Temizle", command=self.tablolari_temizle)
        menu_dosya.add_command(label="Yedeklenmiş Veriler Klasörünü Temizle", command=self.clear_backup_folder)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Hakkında", command=self.hakkinda)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Çıkış", command=self.on_closing)
        menubar.add_cascade(label="Dosya", menu=menu_dosya)

        menu_kalite_kontrol = tk.Menu(menubar, tearoff=0)
        menu_kalite_kontrol.add_command(label="HBTC Formu Oluştur", command=self.hbtc_formu_olustur)
        menu_kalite_kontrol.add_command(label="Verileri Excel'e Aktar", command=lambda: self.excel_e_aktar_sablon(self.tree_kalite, "KaliteKontrol", KALITE_KONTROL_SABLON_EXCEL))
        menubar.add_cascade(label="Kalite Kontrol Menüsü", menu=menu_kalite_kontrol)

        menu_yuzde_sapma = tk.Menu(menubar, tearoff=0)
        menu_yuzde_sapma.add_command(label="Verileri Excel'e Aktar", command=lambda: self.excel_e_aktar_sablon(self.tree_yuzde, "YuzdeSapma", YUZDE_SAPMA_SABLON_EXCEL))
        menubar.add_cascade(label="Yüzde Sapma Menüsü", menu=menu_yuzde_sapma)

        menu_durum_tespiti = tk.Menu(menubar, tearoff=0)
        menu_durum_tespiti.add_command(label="Günü Geçen Ölçümler", command=self.goster_gunu_gecen_olcumler)
        menu_durum_tespiti.add_command(label="Günü Yaklaşan Ölçümler", command=self.goster_gunu_yaklasan_olcumler)
        menubar.add_cascade(label="Durum Tespiti", menu=menu_durum_tespiti)

        menu_hesap = tk.Menu(menubar, tearoff=0)
        menu_hesap.add_command(label="Hesap Makinesi", command=self.open_calculator, accelerator="Ctrl+H")
        menu_hesap.add_command(label="Vücut Kitle İndeksi Hesaplama", command=self.open_bmi_calculation_dialog)
        menubar.add_cascade(label="Hesap Makinesi", menu=menu_hesap)

        menu_takvim = tk.Menu(menubar, tearoff=0)
        menu_takvim.add_command(label="Takvim", command=self.open_calendar, accelerator="Ctrl+T")
        menubar.add_cascade(label="Takvim", menu=menu_takvim)

        self.master.config(menu=menubar)
        self.master.bind('<Control-h>', lambda event: self.open_calculator())
        self.master.bind('<Control-t>', lambda event: self.open_calendar())

    def hbtc_formu_olustur(self):
        if not PYTHON_DOCX_AVAILABLE:
            messagebox.showerror("Kütüphane Hatası",
                                 "Word formu oluşturma için 'python-docx' kütüphanesi kurulu değil.\n"
                                 "Lütfen 'pip install python-docx' komutu ile kurun ve programı yeniden başlatın.",
                                 parent=self.master)
            return

        if not self.tree_kalite.get_children():
            messagebox.showinfo("Veri Yok", "Kalite Kontrol tablosu boş. HBTC formu oluşturulacak veri bulunmuyor.", parent=self.master)
            return

        sablon_yolu = os.path.join(SABLONLAR_DIR, HBTC_TEMPLATE_FILE)
        if not os.path.exists(sablon_yolu):
            messagebox.showerror("Şablon Bulunamadı",
                                 f"Şablon dosyası bulunamadı:\n{sablon_yolu}\n"
                                 f"Lütfen '{HBTC_TEMPLATE_FILE}' dosyasını '{SABLONLAR_DIR}' klasörüne yerleştirin.",
                                 parent=self.master)
            return
        
        self.start_islem_label()
        self.master.update_idletasks() 

        try:
            document = Document(sablon_yolu)
            if not document.tables:
                messagebox.showerror("Şablon Hatası", "Şablon dosyasında doldurulacak tablo bulunamadı.", parent=self.master)
                return

            table = document.tables[0]

            col_map = {
                "Tarih": 1, "Cihaz Tipi - Marka": 2, "Cihaz Seri No": 3,
                "L1": 4, "L2": 5, "L3": 6, "Glukometrenin Geldiği Birim": 7
            }
            s_tarih_idx, s_cihaz_adi_idx, s_seri_no_idx = 0, 1, 2
            s_sonuc_idx, s_min_max_idx, s_bolum_idx = 3, 4, 5
            min_max_degerleri = ["36-108 mg/dl", "144-216 mg/dl", "252-396 mg/dl"]

            def set_cell_text(cell, text_lines, is_multiline=False):
                cell.text = ''
                p = cell.paragraphs[0]
                p.text = ''
                if isinstance(text_lines, str): text_lines = [text_lines]
                for i, line_text in enumerate(text_lines):
                    run = p.add_run(line_text)
                    run.font.name = 'Cambria'
                    run.font.size = Pt(12)
                    if is_multiline and i < len(text_lines) - 1:
                        run.add_break(WD_BREAK.LINE)

            for item_id in self.tree_kalite.get_children():
                values = self.tree_kalite.item(item_id)['values']
                row_cells = table.add_row().cells

                set_cell_text(row_cells[s_tarih_idx], str(values[col_map["Tarih"]]))
                set_cell_text(row_cells[s_cihaz_adi_idx], str(values[col_map["Cihaz Tipi - Marka"]]))
                
                seri_no_ana = str(values[col_map["Cihaz Seri No"]])
                son_4_hane = seri_no_ana[-4:] if len(seri_no_ana) >= 4 else ""
                ana_seri = seri_no_ana[:-4] if len(seri_no_ana) > 4 else seri_no_ana
                set_cell_text(row_cells[s_seri_no_idx], [ana_seri, son_4_hane], is_multiline=True)

                l1, l2, l3 = str(values[col_map["L1"]]), str(values[col_map["L2"]]), str(values[col_map["L3"]])
                sonuc_lines = [f"L1 {l1} mg/dl", f"L2 {l2} mg/dl", f"L3 {l3} mg/dl"]
                set_cell_text(row_cells[s_sonuc_idx], sonuc_lines, is_multiline=True)
                set_cell_text(row_cells[s_min_max_idx], min_max_degerleri, is_multiline=True)
                set_cell_text(row_cells[s_bolum_idx], str(values[col_map["Glukometrenin Geldiği Birim"]]))

            timestamp = datetime.now().strftime('%Y.%m.%d_%H.%M')
            output_filename = f"HBTC_Kalite_Kontrol_Formu_{timestamp}.docx"
            output_path = os.path.join(WORD_FORMS_DIR, output_filename)
            document.save(output_path)

            try:
                if os.name == 'nt': os.startfile(output_path)
                elif sys.platform == 'darwin': subprocess.run(['open', output_path], check=True)
                else: subprocess.run(['xdg-open', output_path], check=True)
            except Exception as e_open:
                print(f"Word dosyası ({output_path}) otomatik olarak açılamadı: {e_open}")

        except Exception as e:
            messagebox.showerror("Form Oluşturma Hatası", f"HBTC formu oluşturulurken bir hata oluştu: {e}", parent=self.master)
        finally:
            self.stop_islem_label()


    def excel_e_aktar_sablon(self, tree, tablo_adi_kisaltmasi, sablon_dosya_adi):
        if not tree.get_children():
            messagebox.showinfo("Veri Yok", f"{tablo_adi_kisaltmasi} tablosunda aktarılacak veri bulunmuyor.", parent=self.master)
            return

        sablon_yolu = os.path.join(SABLONLAR_DIR, sablon_dosya_adi)
        if not os.path.exists(sablon_yolu):
            messagebox.showerror("Şablon Bulunamadı",
                                 f"Excel şablon dosyası bulunamadı:\n{sablon_yolu}\n"
                                 f"Lütfen '{sablon_dosya_adi}' dosyasını '{SABLONLAR_DIR}' klasörüne yerleştirin.",
                                 parent=self.master)
            return
        
        self.start_islem_label()
        self.master.update_idletasks()

        try:
            wb = openpyxl.load_workbook(sablon_yolu)
            ws = wb.active
            tree_headers = tree['columns']
            sablon_headers = [str(cell.value).strip() if cell.value is not None else "" for cell in ws[1]]
            
            start_row_excel = 2
            for row_id in tree.get_children():
                values = tree.item(row_id)['values']
                for col_idx, header_title_tree in enumerate(tree_headers):
                    if header_title_tree == "Bir Sonraki Gelinecek Tarih": # Bu sütunu atla
                        continue
                    try:
                        target_col_excel = sablon_headers.index(header_title_tree) + 1
                        cell_value = values[col_idx]
                        current_cell = ws.cell(row=start_row_excel, column=target_col_excel, value=cell_value)
                        current_cell.alignment = Alignment(horizontal='center', vertical='center') # Ortala

                        if tablo_adi_kisaltmasi == "YuzdeSapma" and header_title_tree == "Yüzde Sapma":
                            try:
                                str_value = str(cell_value).replace('%', '').strip()
                                if str_value:
                                    numeric_value = float(str_value)
                                    if numeric_value > 9.99:
                                        current_cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
                                        current_cell.font = Font(name='Calibri', bold=True, color="FFFFFF")
                                    else:
                                        current_cell.font = Font(name='Calibri')
                                        current_cell.fill = PatternFill(fill_type=None)
                            except ValueError:
                                current_cell.font = Font(name='Calibri')
                                current_cell.fill = PatternFill(fill_type=None)
                        else:
                             current_cell.font = Font(name='Calibri')
                             current_cell.fill = PatternFill(fill_type=None)
                    except ValueError:
                        # print(f"Uyarı: '{header_title_tree}' başlığı şablonda bulunamadı.")
                        continue
                start_row_excel += 1

            if not os.path.exists(EXCEL_OUTPUT_DIR):
                os.makedirs(EXCEL_OUTPUT_DIR)
            timestamp = datetime.now().strftime('%Y.%m.%d_%H.%M')
            output_file_name = f"{tablo_adi_kisaltmasi}_Verileri_{timestamp}.xlsx"
            output_file_path = os.path.join(EXCEL_OUTPUT_DIR, output_file_name)
            wb.save(output_file_path)

            try:
                if os.name == 'nt': os.startfile(output_file_path)
                elif sys.platform == 'darwin': subprocess.run(['open', output_file_path], check=True)
                else: subprocess.run(['xdg-open', output_file_path], check=True)
            except Exception as e_open:
                print(f"Excel dosyası ({output_file_path}) otomatik olarak açılamadı: {e_open}")

        except ImportError:
             messagebox.showerror("Kütüphane Hatası", "Excel'e aktarma için 'openpyxl' kütüphanesi gereklidir.\nLütfen 'pip install openpyxl' komutu ile yükleyin.", parent=self.master)
        except Exception as e:
            messagebox.showerror("Excel'e Aktarma Hatası", f"Veriler Excel'e aktarılırken bir hata oluştu: {e}", parent=self.master)
        finally:
            self.stop_islem_label()


    def open_calculator(self):
        try:
            if os.name == 'nt': subprocess.run("calc.exe", shell=True, check=True)
            elif sys.platform == 'darwin': subprocess.run(["open", "-a", "Calculator"], check=True)
            else:
                found = False
                for calc_cmd in ["gnome-calculator", "kcalc", "xcalc"]:
                    try: subprocess.run([calc_cmd], check=True); found = True; break
                    except FileNotFoundError: continue
                if not found: messagebox.showwarning("Hesap Makinesi", "Standart hesap makinesi komutları bulunamadı.", parent=self.master)
        except (FileNotFoundError, subprocess.CalledProcessError) as e:
            messagebox.showerror("Hata", f"Hesap makinesi açılamadı: {e}", parent=self.master)

    def open_calendar(self):
        top = tk.Toplevel(self.master)
        top.title("Takvim")
        top.iconbitmap(CALENDAR_ICON_PATH)
        top.geometry("400x350") # Biraz daha geniş ve yüksek başlangıç
        #top.resizable(True, True) # Genişletilebilir yapıldı
        top.transient(self.master)
        top.grab_set()

        # Ekran ortasına yerleştirme
        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f"{width}x{height}+{x}+{y}")

        cal = Calendar(top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day,
                      locale="tr_TR", date_pattern="dd.mm.yyyy",
                      font="Arial 10", headersfont="Arial 11 bold")
        cal.pack(padx=15, pady=15, fill="both", expand=True) # Takvim fill="both" ve expand=True ile pencereyle birlikte büyüsün

        top.wait_window()

    def birim_ekle_pencere(self):
        def tamam_click():
            yeni_birim = entry.get().strip().upper()
            if not yeni_birim:
                messagebox.showwarning("Uyarı", "Birim adı boş olamaz!", parent=top)
                return
            if not all(c.isalpha() or c.isspace() or c in "ĞÜŞİÖÇ" for c in yeni_birim if c.strip()):
                messagebox.showwarning("Uyarı", "Birim adı sadece harf, Türkçe karakterler ve boşluk içerebilir!", parent=top)
                return
            if yeni_birim in self.birimler:
                messagebox.showwarning("Uyarı", "Birim zaten mevcut!", parent=top)
                return

            # Veritabanına ekle
            try:
                conn = sqlite3.connect(DB_FILE)
                cursor = conn.cursor()
                cursor.execute("INSERT INTO birimler (birim_adi) VALUES (?)", (yeni_birim,))
                conn.commit()
                conn.close()
                self.load_birimler_from_db() # Verileri yeniden yükle
                self.cmb_birim.set(yeni_birim)
                top.destroy()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Birim eklenirken hata: {e}", parent=top)

        def on_key_release(event):
            current = entry.get()
            entry.delete(0, tk.END)
            entry.insert(0, current.upper())

        top = tk.Toplevel(self.master)
        top.title("Birim Ekle")
        top.geometry("350x130")
        top.resizable(False, False)
        top.transient(self.master)
        top.grab_set()

        if self.plus_icon:
            top.iconbitmap(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Resources", "plus.ico"))

        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f"{width}x{height}+{x}+{y}")

        ttk.Label(top, text="Yeni Birim Adı:").pack(pady=10)
        entry = ttk.Entry(top, width=40)
        entry.pack(pady=5)
        entry.bind("<KeyRelease>", on_key_release)
        ttk.Button(top, text="Tamam", command=tamam_click).pack(pady=15)
        entry.focus_set()
        top.wait_window()

    def birim_sil(self):
        secilen = self.cmb_birim.get()
        if not secilen:
            messagebox.showwarning("Uyarı", "Silinecek birim seçiniz!")
            return
        if messagebox.askokcancel("Onay", f"'{secilen}' birimini silmek istediğinize emin misiniz?"):
            # Veritabanından sil
            try:
                conn = sqlite3.connect(DB_FILE)
                cursor = conn.cursor()
                cursor.execute("DELETE FROM birimler WHERE birim_adi = ?", (secilen,))
                conn.commit()
                conn.close()
                self.load_birimler_from_db() # Verileri yeniden yükle
                if self.birimler:
                    self.cmb_birim.current(0)
                else:
                    self.cmb_birim.set("")
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Birim silinirken hata: {e}")

    def genel_alan_kontrol(self):
        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        son4 = self.cmb_son4hane.get().strip().upper()

        if not tip:
            messagebox.showwarning("Uyarı", "Cihaz Tipi - Marka seçiniz!")
            self.cmb_device_type.focus_set()
            return False
        if not seri:
            messagebox.showwarning("Uyarı", "Cihaz Seri No seçiniz!")
            self.cmb_device_serial.focus_set()
            return False
        if not (son4 and len(son4) == 4 and son4.isalnum()):
            messagebox.showwarning("Uyarı", "Son 4 Hane 4 karakter ve alfanümerik olmalıdır!")
            self.cmb_son4hane.focus_set()
            return False
        if not birim:
            messagebox.showwarning("Uyarı", "Glukometrenin Geldiği Birim seçiniz!")
            self.cmb_birim.focus_set()
            return False

        baska_birim_atanmis = self.check_device_availability(birim, tip, seri, son4)
        if baska_birim_atanmis:
            messagebox.showerror("Kayıt Çakışması",
                               f"Bu Seri Numaralı Cihaz ({seri}{son4}) \n zaten '{baska_birim_atanmis}' birimine kayıtlıdır.\n"
                               "Girdiğiniz Son 4 Haneyi tekrar kontrol ediniz.")
            return False
        return True

    def ayarla_sonraki_tarih(self, olcum_tarihi_str, gun_ekle):
        try:
            if " " in olcum_tarihi_str:
                olcum_tarihi_dt = datetime.strptime(olcum_tarihi_str.split(" ")[0], "%d.%m.%Y").date()
            else:
                olcum_tarihi_dt = datetime.strptime(olcum_tarihi_str, "%d.%m.%Y").date()

            hedef_tarih = olcum_tarihi_dt + timedelta(days=gun_ekle)
            if hedef_tarih.weekday() == 5:
                hedef_tarih -= timedelta(days=1)
            elif hedef_tarih.weekday() == 6:
                hedef_tarih -= timedelta(days=2)
            return hedef_tarih.strftime("%d.%m.%Y")
        except ValueError:
            print(f"Hatalı tarih formatı: {olcum_tarihi_str}. Bugünkü tarih kullanılıyor.")
            bugun = date.today()
            hedef_tarih = bugun + timedelta(days=gun_ekle)
            if hedef_tarih.weekday() == 5:
                hedef_tarih -= timedelta(days=1)
            elif hedef_tarih.weekday() == 6:
                hedef_tarih -= timedelta(days=2)
            return hedef_tarih.strftime("%d.%m.%Y")

    def tabloya_aktar_kalite(self):
        l_values = self.check_l_entries_valid_for_aktar()
        if l_values is None:
            return

        if not self.genel_alan_kontrol():
            return
        l1, l2, l3 = l_values

        tarih_str = datetime.now().strftime("%d.%m.%Y")
        birim = self.cmb_birim.get()
        cihaz_tipi_marka = self.cmb_device_type.get()
        cihaz_seri_no = self.cmb_device_serial.get()
        son_4_hane = self.cmb_son4hane.get().strip().upper()
        cihaz_seri_no_tam = cihaz_seri_no + son_4_hane  # Seri no ve son 4 haneyi birleştir

        if not self.add_or_update_device_assignment(birim, cihaz_tipi_marka, cihaz_seri_no, son_4_hane):
            return

        sonraki_tarih_str = self.ayarla_sonraki_tarih(tarih_str, 15)

        new_item = self.tree_kalite.insert("", "end", values=(
            self.measurement_no_kalite, tarih_str, cihaz_tipi_marka, cihaz_seri_no_tam,  # Birleştirilmiş seri numarasını kullan
            l1, l2, l3, birim, sonraki_tarih_str
        ))
        self.measurement_no_kalite += 1
        self.txt_l1.delete(0, tk.END); self.txt_l1.config(style="TEntry")
        self.txt_l2.delete(0, tk.END); self.txt_l2.config(style="TEntry")
        self.txt_l3.delete(0, tk.END); self.txt_l3.config(style="TEntry")
        self.tree_kalite.selection_set(new_item)
        self.tree_kalite.see(new_item)
        self.update_status_bar()

    def yuzde_sapma_hesapla_ve_aktar(self):
        if not self.genel_alan_kontrol():
            return
        try:
            glukometre_str = self.txt_glukometre_yuzde.get().strip()
            lab_str = self.txt_lab_yuzde.get().strip()
            if not glukometre_str or not lab_str:
                messagebox.showerror("Hata", "Glukometre ve Laboratuvar ölçüm alanları boş olamaz!")
                return
            glukometre = int(float(glukometre_str))
            lab = int(float(lab_str))
            if glukometre == 0 and lab == 0:
                messagebox.showerror("Hata", "Glukometre ve Laboratuvar ölçümleri aynı anda sıfır olamaz!")
                return
            if glukometre == lab:
                yuzde_sapma = 0.0
            elif min(glukometre, lab) == 0:
                yuzde_sapma = 100.0 if max(glukometre, lab) != 0 else 0.0
            else:
                yuzde_sapma = abs((glukometre - lab) / min(glukometre, lab)) * 100

            tarih_str = datetime.now().strftime("%d.%m.%Y")
            birim = self.cmb_birim.get()
            cihaz_tipi_marka = self.cmb_device_type.get()
            cihaz_seri_no = self.cmb_device_serial.get()
            son_4_hane = self.cmb_son4hane.get().strip().upper()
            cihaz_seri_no_tam = cihaz_seri_no + son_4_hane  # Seri no ve son 4 haneyi birleştir

            if not self.add_or_update_device_assignment(birim, cihaz_tipi_marka, cihaz_seri_no, son_4_hane):
                return

            sonraki_tarih_yuzde_str = self.ayarla_sonraki_tarih(tarih_str, 30)
            tags_to_apply = ('high_deviation_tree',) if yuzde_sapma > 9.99 else ()

            new_item = self.tree_yuzde.insert("", "end", values=(
                self.measurement_no_yuzde, tarih_str, cihaz_tipi_marka, cihaz_seri_no_tam,  # Birleştirilmiş seri numarasını kullan
                glukometre, lab, f"{yuzde_sapma:.2f}%", birim, sonraki_tarih_yuzde_str
            ), tags=tags_to_apply)
            self.measurement_no_yuzde += 1
            self.txt_glukometre_yuzde.delete(0, tk.END)
            self.txt_lab_yuzde.delete(0, tk.END)
            self.tree_yuzde.selection_set(new_item)
            self.tree_yuzde.see(new_item)
            self.update_status_bar()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli sayısal değerler girin!")

    def treeview_sort_column(self, tv, col, reverse):
        data_list = []
        for k in tv.get_children(''):
            val = tv.set(k, col)
            data_list.append((val, k))

        try:
            if col == "No":
                data_list.sort(key=lambda t: int(t[0]), reverse=reverse)
            elif col in ["L1", "L2", "L3", "Glukometre Ölçümü", "Lab. Ölçümü"]:
                data_list.sort(key=lambda t: int(float(str(t[0]).replace('%',''))), reverse=reverse)
            elif col == "Yüzde Sapma":
                data_list.sort(key=lambda t: float(str(t[0]).replace('%','')), reverse=reverse)
            elif col in ["Tarih", "Bir Sonraki Gelinecek Tarih"]:
                data_list.sort(key=lambda t: datetime.strptime(str(t[0]), "%d.%m.%Y"), reverse=reverse)
            else:
                data_list.sort(key=lambda t: locale.strxfrm(str(t[0])), reverse=reverse)
        except ValueError as e:
            print(f"Sıralama hatası: {e} - Sütun: {col}, Değer: Unknown")
            data_list.sort(key=lambda t: str(t[0]).upper(), reverse=reverse)
        except NameError:
            data_list.sort(key=lambda t: str(t[0]).upper(), reverse=reverse)

        for index, (_, k) in enumerate(data_list):
            tv.move(k, '', index)
        tv.heading(col, command=lambda: self.treeview_sort_column(tv, col, not reverse))

    def on_double_click(self, event, tree):
        if self.editing_entry:
            self.editing_entry.destroy()
            self.editing_entry = None

        region = tree.identify("region", event.x, event.y)
        if region != "cell":
           
            return
        rowid = tree.identify_row(event.y)
        column_id_str = tree.identify_column(event.x)
        if not rowid or not column_id_str:
            return

        x, y, width, height = tree.bbox(rowid, column_id_str)
        col_index = int(column_id_str.replace("#", "")) - 1

        columns = tree.cget("columns")
        if isinstance(columns, str):
            columns = columns.split()
        if col_index < 0 or col_index >= len(columns):
            return

        column_name = columns[col_index]
        item = tree.item(rowid)
        current_value = item['values'][col_index]

        if column_name == "No":
            messagebox.showwarning("Uyarı", "No sütunu düzenlenemez!")
            return

        self.editing_entry = ttk.Entry(tree)
        self.editing_entry.place(x=x, y=y, width=width, height=height, anchor='nw')
        self.editing_entry.insert(0, current_value)
        self.editing_entry.focus_set()
        self.editing_entry.select_range(0, tk.END)

        def save_edit(event_save=None):
            if not self.editing_entry: return
            new_val = self.editing_entry.get().strip()
            values = list(item['values'])

            if column_name in ["Tarih", "Bir Sonraki Gelinecek Tarih"]:
                try:
                    datetime.strptime(new_val, "%d.%m.%Y")
                except ValueError:
                    messagebox.showerror("Hata", f"Geçersiz tarih formatı: {new_val}. GG.AA.YYYY kullanın.")
                    if self.editing_entry: self.editing_entry.destroy()
                    self.editing_entry = None
                    return

            if tree == self.tree_kalite:
                if column_name == "L1":
                    try: val = int(new_val); assert 36 <= val <= 108
                    except (ValueError, AssertionError): messagebox.showerror("Hata", "L1: 36-108!"); return
                elif column_name == "L2":
                    try: val = int(new_val); assert 144 <= val <= 216
                    except (ValueError, AssertionError): messagebox.showerror("Hata", "L2: 144-216!"); return
                elif column_name == "L3":
                    try: val = int(new_val); assert 252 <= val <= 396
                    except (ValueError, AssertionError): messagebox.showerror("Hata", "L3: 252-396!"); return

            if tree == self.tree_yuzde:
                if column_name in ["Glukometre Ölçümü", "Lab. Ölçümü"]:
                    try: int(new_val)
                    except ValueError: messagebox.showerror("Hata", f"{column_name} tamsayı olmalıdır."); return
                elif column_name == "Yüzde Sapma":
                    try:
                        assert new_val.endswith('%'); float(new_val.replace('%',''))
                    except (AssertionError, ValueError): messagebox.showerror("Hata", "Yüzde Sapma formatı hatalı (örn: 10.50%)."); return

                values[col_index] = new_val
                if column_name in ["Glukometre Ölçümü", "Lab. Ölçümü"]:
                    try:
                        gluk_idx = columns.index("Glukometre Ölçümü")
                        lab_idx = columns.index("Lab. Ölçümü")
                        sapma_idx = columns.index("Yüzde Sapma")
                        glukometre = int(values[gluk_idx])
                        lab = int(values[lab_idx])
                        if glukometre == lab: yuzde_sapma = 0.0
                        elif min(glukometre, lab) == 0: yuzde_sapma = 100.0 if max(glukometre, lab) != 0 else 0.0
                        else: yuzde_sapma = abs((glukometre - lab) / min(glukometre, lab)) * 100
                        values[sapma_idx] = f"{yuzde_sapma:.2f}%"
                        tags = ('high_deviation_tree',) if yuzde_sapma > 9.99 else ()
                        tree.item(rowid, tags=tags)
                    except ValueError: messagebox.showerror("Hata", "Yüzde sapma hesaplanamadı!"); return

            values[col_index] = new_val
            if column_name == "Tarih":
                sonraki_tarih_idx = columns.index("Bir Sonraki Gelinecek Tarih")
                gun_ekle = 15 if tree == self.tree_kalite else 30
                values[sonraki_tarih_idx] = self.ayarla_sonraki_tarih(new_val, gun_ekle)

            tree.item(rowid, values=tuple(values))
            if self.editing_entry: self.editing_entry.destroy()
            self.editing_entry = None

        self.editing_entry.bind("<Return>", save_edit)
        self.editing_entry.bind("<FocusOut>", save_edit)
        self.editing_entry.bind("<Escape>", lambda e: (self.editing_entry.destroy(), setattr(self, 'editing_entry', None)))


    def show_context_menu(self, event, tree):
        rowid = tree.identify_row(event.y)
        if rowid:
            if not tree.selection():
                tree.selection_set(rowid)
            elif rowid not in tree.selection():
                tree.selection_set(rowid)

            if tree == self.tree_kalite:
                self.context_menu_kalite.post(event.x_root, event.y_root)
            elif tree == self.tree_yuzde:
                self.context_menu_yuzde.post(event.x_root, event.y_root)


    def satir_sil(self, tree, is_kalite_table):
        selected_items = tree.selection()
        if not selected_items:
            messagebox.showwarning("Uyarı", "Silmek için bir satır seçiniz!")
            return
        msg = "Seçilen satırı silmek istediğinize emin misiniz?"
        if len(selected_items) > 1:
            msg = f"{len(selected_items)} adet seçili satırı silmek istediğinize emin misiniz?"
        if messagebox.askokcancel("Onay", msg):
            for sel_item in selected_items:
                tree.delete(sel_item)
            self.guncelle_no_sutunu(tree, is_kalite_table)
            self.update_status_bar()

    def guncelle_no_sutunu(self, tree, is_kalite_table):
        all_items = tree.get_children('')
        for idx, item_id in enumerate(all_items, start=1):
            current_values = list(tree.item(item_id, "values"))
            current_values[0] = idx
            tree.item(item_id, values=tuple(current_values))


        if is_kalite_table:
            self.measurement_no_kalite = len(all_items) + 1
        else:
            self.measurement_no_yuzde = len(all_items) + 1

    def tablolari_temizle(self):
        if messagebox.askokcancel("Onay", "Tüm tablolardaki verileri silmek istediğinize emin misiniz? Bu işlem yedeklenmiş verileri etkilemez, sadece mevcut görünümü temizler."):
            for item in self.tree_kalite.get_children():
                self.tree_kalite.delete(item)
            for item in self.tree_yuzde.get_children():
                self.tree_yuzde.delete(item)
            self.measurement_no_kalite = 1
            self.measurement_no_yuzde = 1
            self.tables_cleared_this_session = True
            self.update_status_bar()


    def clear_backup_folder(self):
        msg = "BU İŞLEM TAMAMEN GÜVENLİDİR:\n\n" \
              "- Son 2 yedekleme korunacak, Sadece daha eski yedeklemeler silinecektir \n" \
              "\n" \
              "Devam etmek istiyor musunuz?"

        if not messagebox.askokcancel("Yedeklenmiş Verileri Temizleme", msg, parent=self.master):
            return

        try:

            backup_path = BACKUP_DIR
            if not os.path.exists(backup_path):
                messagebox.showinfo("Bilgi", "Yedeklenmiş Veriler klasörü bulunamadı.", parent=self.master)
                return

            all_files = glob.glob(os.path.join(backup_path, "*.csv"))
            all_files = [f for f in all_files if not os.path.basename(f).startswith('~$')]

            kalite_yedekleri = sorted([f for f in all_files if "Kalite_Kontrol_Olcumleri_Yedek_" in os.path.basename(f)], key=os.path.getctime)
           
            yuzde_yedekleri = sorted([f for f in all_files if "Yuzde_Sapma_Olcumleri_Yedek_" in os.path.basename(f)], key=os.path.getctime)

            files_to_delete = []
            if len(kalite_yedekleri) > 2:
                files_to_delete.extend(kalite_yedekleri[:-2])
            if len(yuzde_yedekleri) > 2:
                files_to_delete.extend(yuzde_yedekleri[:-2])

            if not files_to_delete:
                messagebox.showinfo("Bilgi", "Silinecek daha eski yedek bulunamadı (Her tablo için son 2 yedek dosyası korunuyor).", parent=self.master)
                return

            deleted_count = 0
            for file_path_del in files_to_delete:
                try:
                    os.remove(file_path_del)
                    print(f"Silindi: {os.path.basename(file_path_del)}")
                    deleted_count +=1
                except Exception as e_del:
                    print(f"Silme hatası: {os.path.basename(file_path_del)} - {e_del}")

            if deleted_count > 0:
                messagebox.showinfo("Bilgi", f"Her tablo için son 2 yedekleme saklandı. {deleted_count} adet eski yedek dosyası silindi. ", parent=self.master)
            else:
                messagebox.showinfo("Bilgi", "Silinecek ek yedek bulunamadı.", parent=self.master)

        except Exception as e:
            messagebox.showerror("Hata", f"Yedeklenmis Veriler klasörü temizlenirken bir hata oluştu: {e}", parent=self.master)


    def _tarih_farki_hesapla(self, hedef_tarih_str):
        try:
            hedef_tarih = datetime.strptime(hedef_tarih_str, "%d.%m.%Y").date()
            bugun = date.today()
            fark = (hedef_tarih - bugun).days
            return fark
        except ValueError:
            return None

    def _topla_olcum_verileri(self, durum_tipi):
        bulunan_olcumler = []
        bugun = date.today()
        sira_no = 1

        for item_id in self.tree_kalite.get_children():
            values = self.tree_kalite.item(item_id, 'values')
            if len(values) > 8:
                birim_adi = values[7]
                cihaz_seri_no = values[3]
                gelmesi_gereken_tarih_str = values[8]

                try:
                    gelmesi_gereken_tarih = datetime.strptime(gelmesi_gereken_tarih_str, "%d.%m.%Y").date()
                    fark_gun = (gelmesi_gereken_tarih - bugun).days

                    if durum_tipi == "gecen" and fark_gun < 0:
                        gecen_gun_str = f"{-fark_gun} gün geçti"
                        bulunan_olcumler.append((sira_no, "Kalite Kontrol", birim_adi, cihaz_seri_no, gelmesi_gereken_tarih_str, gecen_gun_str))
                        sira_no +=1
                    elif durum_tipi == "yaklasan" and 0 <= fark_gun <= 2:
                        kalan_gun_str = f"{fark_gun} gün kaldı"
                        if fark_gun == 0: kalan_gun_str = "Bugün"
                        bulunan_olcumler.append((sira_no, "Kalite Kontrol", birim_adi, cihaz_seri_no, gelmesi_gereken_tarih_str, kalan_gun_str))
                        sira_no += 1
                except ValueError:
                    print(f"Kalite tablosunda hatalı tarih formatı: {gelmesi_gereken_tarih_str} - Cihaz: {cihaz_seri_no}")
                    continue

        for item_id in self.tree_yuzde.get_children():
            values = self.tree_yuzde.item(item_id, 'values')
            if len(values) > 8:
                birim_adi = values[7]
                cihaz_seri_no = values[3]
                gelmesi_gereken_tarih_str = values[8]
                try:
                    gelmesi_gereken_tarih = datetime.strptime(gelmesi_gereken_tarih_str, "%d.%m.%Y").date()
                    fark_gun = (gelmesi_gereken_tarih - bugun).days

                    if durum_tipi == "gecen" and fark_gun < 0:
                        gecen_gun_str = f"{-fark_gun} gün geçti"
                        bulunan_olcumler.append((sira_no, "Yüzde Sapma", birim_adi, cihaz_seri_no, gelmesi_gereken_tarih_str, gecen_gun_str))
                        sira_no += 1
                    elif durum_tipi == "yaklasan" and 0 <= fark_gun <= 2:
                        kalan_gun_str = f"{fark_gun} gün kaldı"
                        if fark_gun == 0: kalan_gun_str = "Bugün"
                        bulunan_olcumler.append((sira_no, "Yüzde Sapma", birim_adi, cihaz_seri_no, gelmesi_gereken_tarih_str, kalan_gun_str))
                        sira_no += 1
                except ValueError:
                    print(f"Yüzde Sapma tablosunda hatalı tarih formatı: {gelmesi_gereken_tarih_str} - Cihaz: {cihaz_seri_no}")
                    continue

        return bulunan_olcumler

    def _goster_durum_penceresi(self, baslik, olcum_listesi):
        if not olcum_listesi:
            messagebox.showinfo(baslik, "Bu kriteri karşılayan ölçüm bulunmuyor.", parent=self.master)
            return

        top = tk.Toplevel(self.master)
        top.title(baslik)
        top.geometry("800x400")

        columns = ("No", "Ölçüm Tipi", "Cihazın Geldiği Birim", "Cihaz Seri No", "Gelmesi Gereken Tarih", "Durum")
        tree_durum = ttk.Treeview(top, columns=columns, show="headings")
        tree_durum.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        vsb = ttk.Scrollbar(top, orient="vertical", command=tree_durum.yview)
        vsb.pack(side="right", fill="y", pady=(10,10))
        tree_durum.configure(yscrollcommand=vsb.set)

        widths = [40, 120, 200, 150, 130, 120]
        for i, col_name in enumerate(columns):
            tree_durum.heading(col_name, text=col_name, 
                         command=lambda c=col_name: self.treeview_sort_column(tree_durum, c, False))
            tree_durum.column(col_name, width=widths[i], anchor=tk.W if i in [2,3] else tk.CENTER, minwidth=widths[i])

        for olcum_verisi in olcum_listesi:
            tree_durum.insert("", "end", values=olcum_verisi)

        btn_kapat = ttk.Button(top, text="Kapat", command=top.destroy)
        btn_kapat.pack(pady=10)

        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f'{width}x{height}+{x}+{y}')
        top.focus_set()
        top.wait_window()


    def goster_gunu_gecen_olcumler(self):
        gecen_olcumler = self._topla_olcum_verileri("gecen")
        self._goster_durum_penceresi("Günü Geçen Ölçümler", gecen_olcumler)

    def goster_gunu_yaklasan_olcumler(self):
        yaklasan_olcumler = self._topla_olcum_verileri("yaklasan")
        self._goster_durum_penceresi("Günü Yaklaşan Ölçümler (Son 2 Gün)", yaklasan_olcumler)

    def hakkinda(self):
        about_win = tk.Toplevel(self.master)
        about_win.title("Hakkında - Glukometre Takip Programı")
        about_win.geometry("500x440")
        about_win.resizable(False, False)
        about_win.transient(self.master)
        about_win.grab_set()

        # Pencere ikonu
        try:
            about_win.iconbitmap(APP_ICON_PATH)
        except:
            pass

        # Ana frame
        main_frame = ttk.Frame(about_win)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Logo ve başlık
        logo_frame = ttk.Frame(main_frame)
        logo_frame.pack(fill="x", pady=(0, 15))

        try:
            logo_img = Image.open(APP_ICON_PATH)
            logo_img = logo_img.resize((96, 96), Image.Resampling.LANCZOS)
            logo_photo = ImageTk.PhotoImage(logo_img)
            logo_label = ttk.Label(logo_frame, image=logo_photo)
            logo_label.image = logo_photo
            logo_label.pack(side="left", padx=(0, 15))
        except:
            pass

        title_frame = ttk.Frame(logo_frame)
        title_frame.pack(side="left", fill="y")
        ttk.Label(title_frame, text="Glukometre Takip Programı", 
                 font=("Cambria", 16, "bold")).pack(anchor="w")
        ttk.Label(title_frame, text="Versiyon 2.1", 
                 font=("Cambria", 10)).pack(anchor="w")
        ttk.Label(title_frame, text="© Gökhan ÇOKŞEN", 
                 font=("Cambria", 10)).pack(anchor="w")
        ttk.Label(title_frame, text="2025", 
                 font=("Cambria", 10)).pack(anchor="w")

        # Bilgi metni
        info_frame = ttk.Frame(main_frame)
        info_frame.pack(fill="both", expand=True)

        info_text = tk.Text(info_frame, wrap="word", height=12, 
                           padx=10, pady=10, font=("Tahoma", 10))
        info_text.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(info_frame, orient="vertical", command=info_text.yview)
        scrollbar.pack(side="right", fill="y")
        info_text.config(yscrollcommand=scrollbar.set)
        #info_text.config(state="normal")

        # Bilgi içeriği
        info_content = """
Glukometre Takip Programı, sağlık kuruluşlarında kullanılan glukometre cihazlarının kalite kontrol ve yüzde sapma hesabı ile performans takibini kolaylaştırmak amacıyla geliştirilmiştir.

Öne Çıkan Özellikler:
• Kullanıcı dostu, anlaşılır arayüz
• Kalite kontrol ve yüzde sapma ölçümlerinin kaydı
• Otomatik Veri Yedekleme
• Excel ve Word raporları oluşturma
• Günü Geçen/Yaklaşan Ölçüm Uyarıları

Ek Özellikler:
• Dahili İnternet Radyosu
• Dahili Dijital Saat
• Vücut Kitle İndeksi Hesaplama
• Takvim ve Hesap Makinası erişimi

Teknik Özellikler:
• Python 3.11 ile geliştirildi
• SQLite veritabanı kullanılıyor
• Modern Tkinter arayüzü

Sistem Gereksinimleri:
• Windows 7/10/11
• Python 3.8+
• 4GB RAM
• 100MB boş disk alanı

Geliştirici: Gökhan ÇOKŞEN
İletişim: g.coksen@gmail.com
Son Güncelleme: 19 Mayıs 2025
"""
        info_text.insert("1.0", info_content)
        info_text.config(state="disabled") # Salt okunur yap

        # Tıklanabilir GitHub linki
        hyperlink = tk.Label(main_frame, text="GitHub: https://github.com/gkhncksn/Glukometre_Takip",
                             fg="blue", cursor="hand2", font=("Arial", 10, "underline"))
        hyperlink.pack(pady=(10, 0))
        def open_link(event):
            import webbrowser
            webbrowser.open("https://github.com/biorap/Glukometre_Takip")
        hyperlink.bind("<Button-1>", open_link)

        # Alt bilgi
        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(fill="x", pady=(10, 0))
        ttk.Label(footer_frame, text="Yapay Zeka yardımıyla kodlanmış olup, kaynak kodları GitHub'da yayınlanmıştır.", 
                 font=("Arial", 8)).pack(side="left")
        ttk.Button(footer_frame, text="Kapat", width=10, 
                  command=about_win.destroy).pack(side="right")

        # Pencereyi ortala
        about_win.update_idletasks()
        width = about_win.winfo_width()
        height = about_win.winfo_height()
        x = (about_win.winfo_screenwidth() // 2) - (width // 2)
        y = (about_win.winfo_screenheight() // 2) - (height // 2)
        about_win.geometry(f"{width}x{height}+{x}+{y}")
        about_win.focus_set()
        about_win.wait_window()

    def save_data_to_timestamped_csv(self):
        self.start_islem_label()
        self.master.update_idletasks()
        try:
            timestamp = datetime.now().strftime("%Y.%m.%d_%H.%M")
            kalite_has_data = bool(self.tree_kalite.get_children())
            yuzde_has_data = bool(self.tree_yuzde.get_children())

            if kalite_has_data:
                kalite_filename_ts = os.path.join(BACKUP_DIR, f"Kalite_Kontrol_Olcumleri_Yedek_{timestamp}.csv")
                with open(kalite_filename_ts, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(self.tree_kalite['columns'])
                    for row_id in self.tree_kalite.get_children():
                        writer.writerow(self.tree_kalite.item(row_id)['values'])
                print(f"Kalite Kontrol verileri {os.path.basename(kalite_filename_ts)} dosyasına yedeklendi.")

            if yuzde_has_data:
                yuzde_filename_ts = os.path.join(BACKUP_DIR, f"Yuzde_Sapma_Olcumleri_Yedek_{timestamp}.csv")
                with open(yuzde_filename_ts, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(self.tree_yuzde['columns'])
                    for row_id in self.tree_yuzde.get_children():
                        writer.writerow(self.tree_yuzde.item(row_id)['values'])
                print(f"Yüzde Sapma verileri {os.path.basename(yuzde_filename_ts)} dosyasına yedeklendi.")

            if not kalite_has_data and not yuzde_has_data:
                 print("Kaydedilecek veri bulunmadığı için yedekleme yapılmadı.")
        except Exception as e:
            messagebox.showerror(".CSV Kaydetme Hatası", f"Veriler .CSV dosyasına kaydedilirken bir hata oluştu:\n{e}", parent=self.master)
        finally:
            self.stop_islem_label()


    def show_backup_selection_dialog(self, kalite_files, yuzde_files):
        dialog = tk.Toplevel(self.master)
        dialog.title("Yedek Dosyalarını Seçin")
        dialog.geometry("600x450")
        dialog.transient(self.master)
        dialog.grab_set()
        dialog.resizable(False, False)

        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(expand=True, fill="both")

        ttk.Label(main_frame, text="Yüklemek istediğiniz yedek dosyalarını seçin:", justify=tk.LEFT).pack(pady=(0,10), anchor='w')

        kalite_frame = ttk.LabelFrame(main_frame, text="Kalite Kontrol Yedekleri", padding="10")
        kalite_frame.pack(pady=5, fill="x")
        kalite_inner_frame = ttk.Frame(kalite_frame)
        kalite_inner_frame.pack(fill="both", expand=True)
        kalite_scrollbar = ttk.Scrollbar(kalite_inner_frame, orient=tk.VERTICAL)
        kalite_listbox = tk.Listbox(kalite_inner_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, yscrollcommand=kalite_scrollbar.set)
        kalite_scrollbar.config(command=kalite_listbox.yview)
        kalite_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        kalite_listbox.pack(side=tk.LEFT, fill="both", expand=True)
        kalite_map = {os.path.basename(f): f for f in kalite_files}
        for f_basename in sorted(kalite_map.keys(), reverse=True):
            kalite_listbox.insert(tk.END, f_basename)

        yuzde_frame = ttk.LabelFrame(main_frame, text="Yüzde Sapma Yedekleri", padding="10")
        yuzde_frame.pack(pady=5, fill="x")
        yuzde_inner_frame = ttk.Frame(yuzde_frame)
        yuzde_inner_frame.pack(fill="both", expand=True)
        yuzde_scrollbar = ttk.Scrollbar(yuzde_inner_frame, orient=tk.VERTICAL)
        yuzde_listbox = tk.Listbox(yuzde_inner_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, yscrollcommand=yuzde_scrollbar.set)
        yuzde_scrollbar.config(command=yuzde_listbox.yview)
        yuzde_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        yuzde_listbox.pack(side=tk.LEFT, fill="both", expand=True)
        yuzde_map = {os.path.basename(f): f for f in yuzde_files}
        for f_basename in sorted(yuzde_map.keys(), reverse=True):
            yuzde_listbox.insert(tk.END, f_basename)

        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=15)

        result = {"kalite": None, "yuzde": None}

        def on_ok():
            selected_kalite_indices = kalite_listbox.curselection()
            selected_yuzde_indices = yuzde_listbox.curselection()

            result["kalite"] = [kalite_map[kalite_listbox.get(i)] for i in selected_kalite_indices]
            result["yuzde"] = [yuzde_map[yuzde_listbox.get(i)] for i in selected_yuzde_indices]

            dialog.destroy()

        def on_cancel():
            result["kalite"] = None
            result["yuzde"] = None
            dialog.destroy()

        ok_button = ttk.Button(button_frame, text="Seçilenleri Yükle", command = on_ok)
        ok_button.pack(side=tk.LEFT, padx=10)
        cancel_button = ttk.Button(button_frame, text="İptal Et / Boş Başlat", command=on_cancel)
        cancel_button.pack(side=tk.LEFT, padx=10)

        dialog.protocol("WM_DELETE_WINDOW", on_cancel)
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')

        dialog.wait_window()

        if result["kalite"] is not None and result["yuzde"] is not None:
            return result["kalite"], result["yuzde"]
        else:
             return None


    def load_data_from_csv(self):
        self.start_islem_label()
        self.master.update_idletasks()
        try:
            load_kalite_files, load_yuzde_files = [], []
            all_kalite_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Kalite_Kontrol_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
            if all_kalite_files: load_kalite_files = [max(all_kalite_files, key=os.path.getctime)]
            all_yuzde_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Yuzde_Sapma_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
            if all_yuzde_files: load_yuzde_files = [max(all_yuzde_files, key=os.path.getctime)]

            loaded_kalite, loaded_yuzde = False, False
            for item in self.tree_kalite.get_children(): self.tree_kalite.delete(item)
            for item in self.tree_yuzde.get_children(): self.tree_yuzde.delete(item)
            self.measurement_no_kalite, self.measurement_no_yuzde = 1, 1

            expected_kalite_headers = self.tree_kalite['columns']
            for file_path in load_kalite_files:
                try:
                    with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                        reader = csv.reader(f, delimiter=';')
                        try: headers = next(reader)
                        except StopIteration: continue
                        if tuple(headers) == expected_kalite_headers:
                            file_has_data = False
                            for row in reader:
                                if len(row) == len(expected_kalite_headers): self.tree_kalite.insert("", "end", values=row); file_has_data = True
                            if file_has_data: loaded_kalite = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

            expected_yuzde_headers = self.tree_yuzde['columns']
            sapma_str_idx = list(expected_yuzde_headers).index("Yüzde Sapma") if "Yüzde Sapma" in expected_yuzde_headers else -1
            for file_path in load_yuzde_files:
                try:
                    with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                        reader = csv.reader(f, delimiter=';')
                        try: headers = next(reader)
                        except StopIteration: continue
                        if tuple(headers) == expected_yuzde_headers:
                            file_has_data = False
                            for i, row in enumerate(reader, start=1):
                                if len(row) == len(expected_yuzde_headers):
                                    tags_to_apply = ()
                                    if sapma_str_idx != -1:
                                        try:
                                            sapma_val = float(row[sapma_str_idx].replace('%','').strip())
                                            if sapma_val > 9.99: tags_to_apply = ('high_deviation_tree',)
                                        except: pass # Ignore tagging errors for malformed data
                                    self.tree_yuzde.insert("", "end", values=row, tags=tags_to_apply)
                                    file_has_data = True
                            if file_has_data: loaded_yuzde = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

            if loaded_kalite: self.guncelle_no_sutunu(self.tree_kalite, True)
            if loaded_yuzde: self.guncelle_no_sutunu(self.tree_yuzde, False)
        finally:
            self.stop_islem_label()
            self.update_status_bar()

    def manuel_yedek_yukle(self):
        if not messagebox.askokcancel("Onay", "Mevcut tablo verileri silinecek ve seçilen yedekler yüklenecektir.\nDevam etmek istiyor musunuz?", parent=self.master):
            return

        self.start_islem_label()
        self.master.update_idletasks()
        try:
            kalite_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Kalite_Kontrol_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
            yuzde_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Yuzde_Sapma_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]

            if not kalite_files and not yuzde_files:
                messagebox.showinfo("Bilgi", "Yüklenecek (boş olmayan) yedek dosya bulunamadı.", parent=self.master)
                return

            result = self.show_backup_selection_dialog(kalite_files, yuzde_files)
            if result:
                load_kalite_files, load_yuzde_files = result
                if not load_kalite_files and not load_yuzde_files: return

                loaded_kalite, loaded_yuzde = False, False
                for item in self.tree_kalite.get_children(): self.tree_kalite.delete(item)
                for item in self.tree_yuzde.get_children(): self.tree_yuzde.delete(item)
                self.measurement_no_kalite, self.measurement_no_yuzde = 1, 1

                expected_kalite_headers = self.tree_kalite['columns']
                for file_path in load_kalite_files:
                    try:
                        if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                            with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                                reader = csv.reader(f, delimiter=';'); headers = next(reader, None)
                                if headers and tuple(headers) == expected_kalite_headers:
                                    file_has_data = False
                                    for row in reader:
                                        if len(row) == len(expected_kalite_headers): self.tree_kalite.insert("", "end", values=row); file_has_data = True
                                    if file_has_data: loaded_kalite = True
                    except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

                expected_yuzde_headers = self.tree_yuzde['columns']
                sapma_str_idx = list(expected_yuzde_headers).index("Yüzde Sapma") if "Yüzde Sapma" in expected_yuzde_headers else -1
                for file_path in load_yuzde_files:
                    try:
                        if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                            with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                                reader = csv.reader(f, delimiter=';'); headers = next(reader, None)
                                if headers and tuple(headers) == expected_yuzde_headers:
                                    file_has_data = False
                                    for i, row in enumerate(reader, start=1):
                                        if len(row) == len(expected_yuzde_headers):
                                            tags_to_apply = ()
                                            if sapma_str_idx != -1:
                                                try:
                                                    sapma_val = float(row[sapma_str_idx].replace('%','').strip())
                                                    if sapma_val > 9.99: tags_to_apply = ('high_deviation_tree',)
                                                except: pass
                                            self.tree_yuzde.insert("", "end", values=row, tags=tags_to_apply)
                                            file_has_data = True
                                    if file_has_data: loaded_yuzde = True
                    except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)
                
                if loaded_kalite: self.guncelle_no_sutunu(self.tree_kalite, True)
                if loaded_yuzde: self.guncelle_no_sutunu(self.tree_yuzde, False)
                self.update_status_bar()
        finally:
            self.stop_islem_label()

    def on_closing(self):
        last_station_name = self.cmb_radyo.get()
        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()
        else:
            self.ffplay_process_id = None # Eğer ffplay işlemi zaten durmuşsa, PID'yi sıfırla
        if last_station_name:
            self.save_setting("last_radio_station", last_station_name)
        self.save_setting("last_radio_volume", str(self.radio_volume.get()))

        kalite_has_data = bool(self.tree_kalite.get_children())
        yuzde_has_data = bool(self.tree_yuzde.get_children())
        tables_have_data = kalite_has_data or yuzde_has_data

        if not tables_have_data and self.tables_cleared_this_session:
            if messagebox.askokcancel("Çıkış Onayı", "Tablolar temizlendi ve yeni veri girilmedi.\nÇıkmak istediğinize emin misiniz?", parent=self.master):
                self.master.destroy()
            return
        elif not tables_have_data and not self.tables_cleared_this_session:
            self.master.destroy()
            return
        else:
            if messagebox.askokcancel("Çıkış ve Yedekle", "Programdan çıkmak ve verileri yeni zaman damgalı dosyalara yedeklemek istiyor musunuz?", parent=self.master):
                self.save_data_to_timestamped_csv() 
                self.master.destroy()

    def load_radio_stations(self):
        RADIO_FILE = "RadioStationsFFMPEG.txt"
        names, name_url_map = [], {}
        if os.path.exists(RADIO_FILE):
            with open(RADIO_FILE, "r", encoding="utf-8-sig") as f:
                for line in f:
                    line = line.strip()
                    if not line or "|" not in line: continue
                    try:
                        name, url = line.split("|", 1)
                        name, url = name.strip(), url.strip()
                        if name and url: names.append(name); name_url_map[name] = url
                    except ValueError: print(f"RadioStationsFFMPEG.txt dosyasında hatalı satır: {line}")
        try: names = sorted(names, key=locale.strxfrm)
        except: names = sorted(names)
        return names, name_url_map

    def play_radio_command(self):
        self.play_radio() # Calls play_radio which will use the current volume

    def play_radio(self, volume_level=None):
        selected_name = self.cmb_radyo.get()
        if not selected_name or selected_name not in self.radio_station_map:
            messagebox.showwarning("Radyo", "Lütfen bir radyo istasyonu seçin!", parent=self.master)
            return
        url = self.radio_station_map[selected_name]

        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()

        actual_volume = volume_level if volume_level is not None else self.radio_volume.get()

        try:
            # Önce programın çalıştığı dizindeki ffmpeg klasöründe ffplay.exe'yi ara
            script_dir = os.path.dirname(os.path.abspath(__file__))
            ffmpeg_dir = os.path.join(script_dir, "ffmpeg")
            ffplay_path = os.path.join(ffmpeg_dir, "ffplay.exe")
            
            if not os.path.exists(ffplay_path):
                # Eğer yerel klasörde yoksa, sistem PATH'ında ara
                ffplay_path = "ffplay"

            startupinfo = None; creation_flags = 0
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = subprocess.SW_HIDE
                creation_flags = subprocess.CREATE_NO_WINDOW
        
            cmd = [ffplay_path, "-nodisp", "-autoexit", "-loglevel", "error", "-volume", str(actual_volume), url]
            self.radio_process = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                                                  startupinfo=startupinfo, creationflags=creation_flags)
            self.update_radio_button_states(playing=True)
        except FileNotFoundError:
            messagebox.showerror("Radyo", 
                           "ffplay bulunamadı!\n"
                           f"Lütfen ffmpeg klasörünü program dizinine ({ffmpeg_dir}) kopyalayın\n"
                           "veya sistem PATH'ına ffmpeg ekleyin.", 
                           parent=self.master)
            self.radio_process = None; self.update_radio_button_states(playing=False)
        except Exception as e:
            messagebox.showerror("Radyo", f"Radyo başlatılamadı: {e}", parent=self.master)
            self.radio_process = None; self.update_radio_button_states(playing=False)

    def stop_radio(self):
        if self.radio_process and self.radio_process.poll() is None:
            try: self.radio_process.terminate()
            except Exception as e: print(f"Radyo durdurulurken hata: {e}")
            finally: self.radio_process = None
        self.update_radio_button_states(playing=False)

    def open_bmi_calculation_dialog(self):
        dialog = tk.Toplevel(self.master)
        dialog.title("Vücut Kitle İndeksi Hesaplama")
        dialog.geometry("300x170")
        dialog.resizable(False, False)
        dialog.transient(self.master)
        dialog.grab_set()

        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f"{width}x{height}+{x}+{y}")

        ttk.Label(dialog, text="Kilonuzu Giriniz (Kg):").pack(pady=5)
        weight_entry = ttk.Entry(dialog)
        weight_entry.pack(pady=5)

        ttk.Label(dialog, text="Boyunuzu Giriniz (Cm):").pack(pady=5)
        height_entry = ttk.Entry(dialog)
        height_entry.pack(pady=5)

        def calculate_bmi():
            try:
                kilo = float(weight_entry.get())
                boy = float(height_entry.get())
                if kilo <= 0 or boy <= 0:
                    messagebox.showerror("Hata", "Kilo ve boy pozitif değerler olmalıdır!", parent=dialog)
                    return
                vki = kilo / ((boy / 100) * (boy / 100))
                durum = ""
                if vki < 18.5: durum = "Zayıf"
                elif 18.5 <= vki < 24.9: durum = "Normal"
                elif 25 <= vki < 29.9: durum = "Fazla Kilolu"
                elif 30 <= vki < 34.9: durum = "Obez (Sınıf 1)"
                elif 35 <= vki < 39.9: durum = "Obez (Sınıf 2)"
                else: durum = "Aşırı Obez (Sınıf 3)"
                messagebox.showinfo("Vücut Kitle İndeksi", f"Vücut Kitle İndeksiniz: {vki:.2f}\nMevcut Durumunuz: {durum}", parent=dialog)
            except ValueError:
                messagebox.showerror("Hata", "Lütfen geçerli sayısal değerler girin!", parent=dialog)

        ttk.Button(dialog, text="Hesapla", command=calculate_bmi).pack(pady=10)
        dialog.wait_window()


if __name__ == '__main__':
    root = tk.Tk()
    app = MainWindow(root)
    root.mainloop()
